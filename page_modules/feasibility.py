import streamlit as st  # type: ignore
import streamlit.components.v1 as components  # type: ignore
import base64
from io import BytesIO

from ui_helpers import _h, _safe_filename, section_header, page_title, celebrate
from analytics import (
    log_event,
    EVENT_GENERATE_FEAS, EVENT_DOWNLOAD_FEAS,
)


def render_feasibility():
    # Clear any previously generated doc when the page is freshly visited
    # so stale downloads from earlier sessions don't persist.
    if st.session_state.get("analytics_last_page") != "feasibility":
        st.session_state.pop("feas_doc", None)
        st.session_state.pop("feas_name", None)

    page_title(
        "Feasibility Assessment",
        "Evaluate crawl feasibility before project kickoff — exports a Word document for the team."
    )

    # Stats strip
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("""
        <div style="background:white;border-radius:12px;padding:18px 20px;border-left:4px solid #1f2937;box-shadow:0 2px 8px rgba(0,0,0,0.07);transition:box-shadow 0.2s;">
            <div style="font-size:0.7rem;color:#94a3b8;text-transform:uppercase;letter-spacing:0.09em;font-weight:700;font-family:'Inter',sans-serif;">Purpose</div>
            <div style="font-size:0.9rem;font-weight:600;color:#111827;margin-top:5px;font-family:'Inter',sans-serif;">Pre-project scoping</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown("""
        <div style="background:white;border-radius:12px;padding:18px 20px;border-left:4px solid #1f2937;box-shadow:0 2px 8px rgba(0,0,0,0.07);transition:box-shadow 0.2s;">
            <div style="font-size:0.7rem;color:#94a3b8;text-transform:uppercase;letter-spacing:0.09em;font-weight:700;font-family:'Inter',sans-serif;">Output</div>
            <div style="font-size:0.9rem;font-weight:600;color:#111827;margin-top:5px;font-family:'Inter',sans-serif;">Word Document (.docx)</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown("""
        <div style="background:white;border-radius:12px;padding:18px 20px;border-left:4px solid #1f2937;box-shadow:0 2px 8px rgba(0,0,0,0.07);transition:box-shadow 0.2s;">
            <div style="font-size:0.7rem;color:#94a3b8;text-transform:uppercase;letter-spacing:0.09em;font-weight:700;font-family:'Inter',sans-serif;">Use Case</div>
            <div style="font-size:0.9rem;font-weight:600;color:#111827;margin-top:5px;font-family:'Inter',sans-serif;">Share with tech / ops team</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    col, side = st.columns([3, 1])
    with side:
        st.markdown("""
        <div style="background:linear-gradient(135deg,#f8fafc 0%,#f1f5f9 100%);
                    border:1px solid #e5e7eb;border-radius:12px;padding:16px;
                    font-family:'Inter',sans-serif;margin-top:8px;">
            <div style="font-size:0.72rem;font-weight:700;color:#6b7280;text-transform:uppercase;
                        letter-spacing:0.08em;margin-bottom:10px;">📄 What's included</div>
            <div style="display:flex;flex-direction:column;gap:7px;">
                <div style="display:flex;align-items:center;gap:8px;font-size:0.78rem;color:#374151;">
                    <span style="color:#16a34a;font-weight:700;">✓</span> Client &amp; requestor info
                </div>
                <div style="display:flex;align-items:center;gap:8px;font-size:0.78rem;color:#374151;">
                    <span style="color:#16a34a;font-weight:700;">✓</span> Full domain list
                </div>
                <div style="display:flex;align-items:center;gap:8px;font-size:0.78rem;color:#374151;">
                    <span style="color:#16a34a;font-weight:700;">✓</span> Crawl type &amp; features
                </div>
                <div style="display:flex;align-items:center;gap:8px;font-size:0.78rem;color:#374151;">
                    <span style="color:#16a34a;font-weight:700;">✓</span> Zipcode requirements
                </div>
                <div style="display:flex;align-items:center;gap:8px;font-size:0.78rem;color:#374151;">
                    <span style="color:#16a34a;font-weight:700;">✓</span> Additional notes
                </div>
            </div>
            <div style="margin-top:12px;padding-top:10px;border-top:1px solid #e5e7eb;
                        font-size:0.72rem;color:#9ca3af;">
                Output: <strong style="color:#374151;">.docx</strong> — share directly with the tech / ops team.
            </div>
        </div>""", unsafe_allow_html=True)
    with col:

        # Requirement Information
        section_header("📌", "Requirement Information")
        c1, c2 = st.columns(2)
        with c1:
            client_name = st.text_input("Client Name", placeholder="e.g., Unilever India", key="feas_client")
        with c2:
            requestor_name = st.text_input("Requestor Name", placeholder="Your name", key="feas_requestor")

        # Domains
        section_header("🌐", "Domain List")
        num_domains = st.number_input("Number of Domains", min_value=1, max_value=50, step=1, value=1, key="feas_num_domains")
        domains = []
        if num_domains <= 6:
            cols = st.columns(min(int(num_domains), 3))
            for i in range(int(num_domains)):
                with cols[i % 3]:
                    d = st.text_input(f"Domain {i+1}", placeholder="example.com", key=f"feas_domain_{i}")
                    if d:
                        domains.append(d)
        else:
            for i in range(int(num_domains)):
                d = st.text_input(f"Domain {i+1}", placeholder="example.com", key=f"feas_domain_{i}")
                if d:
                    domains.append(d)

        # Crawl Configuration
        section_header("⚙️", "Crawl Configuration")
        crawl_options = st.multiselect(
            "Select crawl type and special requirements",
            ["Category Based", "Product URL Input Based", "SOS", "Reviews",
             "Festive Sales Day Crawl", "Banner Crawl", "Others"],
            key="feas_crawl_options",
        )
        crawl_type = None
        if "Category Based" in crawl_options:
            crawl_type = "Category Based"
        elif "Product URL Input Based" in crawl_options:
            crawl_type = "Product URL Input Based"

        crawl_features = [o for o in crawl_options if o not in ["Category Based", "Product URL Input Based"]]
        others_desc = ""
        if "Others" in crawl_features:
            others_desc = st.text_input("If Others, please specify", key="feas_others")

        # Zipcode
        section_header("📍", "Zipcode Requirement")
        zipcode_type = st.radio(
            "Zipcode handling", ["With Zipcode", "Without Zipcode", "Both"],
            horizontal=True, key="feas_zipcode"
        )
        target_city = target_state = target_country = ""
        if zipcode_type in ["With Zipcode", "Both"]:
            st.markdown("**Target Location**")
            c1, c2, c3 = st.columns(3)
            with c1:
                target_city = st.text_input("City", key="feas_city")
            with c2:
                target_state = st.text_input("State", key="feas_state")
            with c3:
                target_country = st.text_input("Country", key="feas_country")

        # Additional
        section_header("📝", "Additional Information")
        additional_notes = st.text_area("Additional details / notes", key="feas_notes", height=120)

        st.markdown("<br>", unsafe_allow_html=True)

        # Generate + Download feasibility document (single button)
        if st.button("📄  Generate & Download Feasibility Document", type="primary", width="stretch"):
            if not client_name:
                st.error("Enter a Client Name before generating the document.")
                st.stop()
            log_event(EVENT_GENERATE_FEAS, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "feasibility", {"client": client_name})
            try:
                with st.spinner("Building document…"):
                    from docx import Document  # type: ignore
                    doc = Document()
                    doc.add_heading(f"{client_name} — Feasibility Requirement", level=1)

                    doc.add_heading("Requirement Information", level=2)
                    doc.add_paragraph(f"Client Name: {client_name}")
                    doc.add_paragraph(f"Requestor Name: {requestor_name}")

                    doc.add_heading("Domains", level=2)
                    for d in domains:
                        doc.add_paragraph(d, style="List Bullet")

                    doc.add_heading("Crawl Configuration", level=2)
                    doc.add_paragraph(f"Crawl Type: {crawl_type or 'Not specified'}")
                    doc.add_paragraph(f"Special Requirements: {', '.join(crawl_features) or 'None'}")
                    if others_desc:
                        doc.add_paragraph(f"Others: {others_desc}")

                    doc.add_heading("Zipcode Requirement", level=2)
                    doc.add_paragraph(f"Zipcode Handling: {zipcode_type}")
                    if zipcode_type in ["With Zipcode", "Both"]:
                        doc.add_heading("Target Location", level=2)
                        doc.add_paragraph(f"City: {target_city}")
                        doc.add_paragraph(f"State: {target_state}")
                        doc.add_paragraph(f"Country: {target_country}")

                    doc.add_heading("Additional Notes", level=2)
                    doc.add_paragraph(additional_notes or "None")

                    buf = BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    _feas_bytes = buf.getvalue()

                log_event(EVENT_DOWNLOAD_FEAS, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "feasibility")
                celebrate(message="Downloading document…", sub=f"{_h(client_name)} feasibility doc is downloading.")
                _feas_b64  = base64.b64encode(_feas_bytes).decode()
                _feas_name = _safe_filename(client_name, "_Feasibility_Requirement.docx")
                _feas_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                components.html(f"""<script>
                (function(){{
                    var b = atob("{_feas_b64}");
                    var a = new Uint8Array(b.length);
                    for(var i=0;i<b.length;i++) a[i]=b.charCodeAt(i);
                    var blob = new Blob([a],{{type:"{_feas_mime}"}});
                    var url  = URL.createObjectURL(blob);
                    var el   = window.parent.document.createElement("a");
                    el.href  = url; el.download = "{_feas_name}";
                    window.parent.document.body.appendChild(el);
                    el.click();
                    window.parent.document.body.removeChild(el);
                    URL.revokeObjectURL(url);
                }})();
                </script>""", height=0)
            except Exception as e:
                st.error(f"Document generation failed — please try again or contact your admin. ({type(e).__name__})")
