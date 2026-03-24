import streamlit as st # type: ignore
import streamlit.components.v1 as components # type: ignore
from io import BytesIO
from datetime import date, datetime, timedelta, timezone
import os
import csv
import base64
import html as _html_mod
import re
import time
import json
import uuid
from pathlib import Path
import pandas as pd # type: ignore
from credentials import verify_password, get_user, MAX_ATTEMPTS, LOCKOUT_SECONDS
from analytics import (
    log_event,
    get_summary,
    load_events,
    PAGE_LABELS,
    EVENT_LABELS,
    EVENT_LOGIN, EVENT_LOGOUT, EVENT_PAGE_VIEW,
    EVENT_GENERATE_REQ_PDF, EVENT_DOWNLOAD_REQ_PDF,
    EVENT_GENERATE_FEAS, EVENT_DOWNLOAD_FEAS,
    EVENT_DOWNLOAD_COST_PDF, EVENT_DOWNLOAD_COST_CSV,
)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────

try:
    from PIL import Image as _PIL_Image  # type: ignore
    _page_icon = _PIL_Image.open("42slogo_top.png") if os.path.exists("42slogo_top.png") else "🔍"
except Exception:
    _page_icon = "🔍"

st.set_page_config(
    page_title="42Signals | Requirement Handling",
    page_icon=_page_icon,
    layout="wide",
    initial_sidebar_state="expanded",
)

LOGO_PATH = str(Path(__file__).parent / "42slogo.png")

# Bidirectional component for secure localStorage-based session persistence.
_session_mgr = components.declare_component(
    "session_manager",
    path=str(Path(__file__).parent / "session_component"),
)

# D3.js bundled locally so mind maps work on servers without internet access.
# Falls back to CDN if the file is missing (dev convenience only).
@st.cache_resource
def _load_d3_inline():
    _path = Path("d3.v7.min.js")
    if _path.exists():
        return f"<script>{_path.read_text()}</script>"
    return '<script src="https://cdn.jsdelivr.net/npm/d3@7/dist/d3.min.js"></script>'

_D3_INLINE = _load_d3_inline()

# ─────────────────────────────────────────────────────────────────────────────
# SECURITY HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _h(value) -> str:
    """HTML-escape any user-supplied value before injecting into unsafe_allow_html contexts.
    Prevents XSS when user text is embedded in st.markdown HTML strings."""
    return _html_mod.escape(str(value), quote=True)


def _safe_filename(name: str, suffix: str = "") -> str:
    """Sanitize a user-supplied string for use as a download filename.
    Strips path-traversal characters, null bytes, and limits length."""
    safe = re.sub(r'[^\w\s\-]', '', str(name), flags=re.UNICODE).strip()
    safe = re.sub(r'\s+', '_', safe)
    safe = safe[:80]  # cap at 80 chars to avoid filesystem limits
    return (safe or "document") + suffix


# ─────────────────────────────────────────────────────────────────────────────
# PERSISTENT SESSION  (file-based, stdlib only, 7-day expiry)
# ─────────────────────────────────────────────────────────────────────────────

_SESSION_FILE = Path(".42s_session.json")
_SESSION_TTL_DAYS = 7

# ── Server-side lockout (per username, survives new tabs / refreshes) ─────────
_LOCKOUT_FILE = Path(".42s_lockout.json")


def _get_lockout(username: str) -> tuple:
    """Return (attempts, lockout_until_timestamp) for a username."""
    try:
        if not _LOCKOUT_FILE.exists():
            return 0, 0.0
        data = json.loads(_LOCKOUT_FILE.read_text())
        rec = data.get(username.strip().lower(), {})
        return int(rec.get("attempts", 0)), float(rec.get("lockout_until", 0.0))
    except (OSError, KeyError, ValueError):
        return 0, 0.0


def _set_lockout(username: str, attempts: int, lockout_until: float) -> None:
    try:
        data = {}
        if _LOCKOUT_FILE.exists():
            try:
                data = json.loads(_LOCKOUT_FILE.read_text())
            except (ValueError, OSError):
                data = {}
        data[username.strip().lower()] = {"attempts": attempts, "lockout_until": lockout_until}
        # Prune entries whose lockout has expired and attempt count is reset
        now = time.time()
        data = {u: v for u, v in data.items() if v.get("lockout_until", 0) > now or v.get("attempts", 0) > 0}
        _LOCKOUT_FILE.write_text(json.dumps(data))
    except OSError:
        pass


def _clear_lockout(username: str) -> None:
    try:
        if not _LOCKOUT_FILE.exists():
            return
        data = json.loads(_LOCKOUT_FILE.read_text())
        data.pop(username.strip().lower(), None)
        _LOCKOUT_FILE.write_text(json.dumps(data))
    except (OSError, ValueError):
        pass


def _save_session(username: str, display_name: str) -> str:
    token = str(uuid.uuid4())
    data = {
        "token": token,
        "username": username,
        "display_name": display_name,
        "expires": (datetime.now(timezone.utc) + timedelta(days=_SESSION_TTL_DAYS)).isoformat(),
    }
    try:
        _SESSION_FILE.write_text(json.dumps(data))
    except OSError:
        pass
    return token


def _load_session(token: str):
    """Return (username, display_name) if token matches a valid non-expired session, else (None, None)."""
    if not token:
        return None, None
    try:
        if not _SESSION_FILE.exists():
            return None, None
        data = json.loads(_SESSION_FILE.read_text())
        if data.get("token") != token:
            return None, None
        if datetime.now(timezone.utc) > datetime.fromisoformat(data["expires"]):
            _SESSION_FILE.unlink(missing_ok=True)
            return None, None
        return data["username"], data["display_name"]
    except (OSError, KeyError, ValueError, json.JSONDecodeError):
        return None, None




def _clear_session() -> None:
    try:
        _SESSION_FILE.unlink(missing_ok=True)
    except OSError:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL CSS
# ─────────────────────────────────────────────────────────────────────────────

# Load Inter font via <link> (faster than @import inside <style>)
st.markdown(
    '<link rel="preconnect" href="https://fonts.googleapis.com">'
    '<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>'
    '<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">',
    unsafe_allow_html=True,
)

st.markdown("""
<style>

/* ── Streamlit chrome ── */
#MainMenu { visibility: hidden; }
footer     { visibility: hidden; }
[data-testid="stToolbar"]    { visibility: hidden; }
[data-testid="stDecoration"] { visibility: hidden; }
[data-testid="stStatusWidget"] { visibility: hidden; }
header     { visibility: hidden; }
header button,
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapsedControl"] {
    visibility: visible !important;
    display: flex !important;
    opacity: 1 !important;
    pointer-events: auto !important;
}

/* ── App background & global font ── */
.stApp {
    background-color: #f0f2f6;
    color: #1f2937 !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
}

/* ── Global text colour (markdown, labels, captions) ── */
.stMarkdown p,
.stMarkdown li,
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3,
.stMarkdown h4, .stMarkdown h5,
.element-container p,
[data-testid="stMainBlockContainer"] > div p,
[data-testid="stWidgetLabel"] p,
[data-testid="stWidgetLabel"] span,
.stCaption, .stText {
    color: #1f2937 !important;
}

/* ── Main block container spacing ── */
.block-container {
    padding-top: 1.8rem !important;
    padding-bottom: 2.5rem !important;
}

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: #ffffff;
    border-right: 1px solid #d1d5db;
    box-shadow: 2px 0 12px rgba(0,0,0,0.05);
}
section[data-testid="stSidebar"] .stMarkdown p,
section[data-testid="stSidebar"] .element-container p,
section[data-testid="stSidebar"] small,
section[data-testid="stSidebar"] .stCaption {
    color: #6b7280 !important;
}
section[data-testid="stSidebar"] hr {
    border-color: #e5e7eb !important;
    margin: 12px 0 !important;
}

/* ── Sidebar nav buttons ── */
section[data-testid="stSidebar"] .stButton > button {
    background: transparent !important;
    border: 1px solid transparent !important;
    border-radius: 8px !important;
    padding: 9px 12px !important;
    font-size: 0.83rem !important;
    font-weight: 500 !important;
    color: #4b5563 !important;
    text-align: left !important;
    justify-content: flex-start !important;
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    box-shadow: none !important;
    letter-spacing: 0 !important;
    transition: background .15s ease, color .15s ease !important;
    margin-bottom: 1px !important;
}
section[data-testid="stSidebar"] .stButton > button:hover {
    background: #f1f5f9 !important;
    color: #1f2937 !important;
    border-color: transparent !important;
    box-shadow: none !important;
    transform: none !important;
}
section[data-testid="stSidebar"] .stButton > button:active {
    background: #e8ecf0 !important;
    box-shadow: none !important;
    transform: none !important;
}
/* Sidebar button inner <p>/<span> — let them inherit the button's colour
   instead of being forced to #1f2937 by the global .element-container p rule */
section[data-testid="stSidebar"] .stButton > button p,
section[data-testid="stSidebar"] .stButton > button span {
    color: inherit !important;
}
/* Use focus-visible so keyboard users see a clear outline; mouse clicks don't trigger it */
section[data-testid="stSidebar"] .stButton > button:focus-visible {
    outline: 2px solid #374151 !important;
    outline-offset: 2px !important;
    box-shadow: none !important;
}
section[data-testid="stSidebar"] .stButton > button:focus:not(:focus-visible) {
    outline: none !important;
    box-shadow: none !important;
}

/* ── Sidebar expander ── */
section[data-testid="stSidebar"] details summary p {
    color: #374151 !important;
    font-weight: 600 !important;
}

/* ── Text inputs & textareas ── */
/* Use broad selectors to cover Streamlit's nested data-baseweb wrappers */
.stTextInput input,
.stTextArea textarea {
    border-radius: 8px !important;
    border: 1.5px solid #c9cfd8 !important;
    background: #ffffff !important;
    font-size: 0.875rem !important;
    color: #1f2937 !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
    padding: 8px 12px !important;
}
.stTextInput input:focus,
.stTextArea textarea:focus {
    border-color: #374151 !important;
    box-shadow: 0 0 0 3px rgba(55,65,81,0.12) !important;
    outline: none !important;
}
.stTextInput input::placeholder,
.stTextArea textarea::placeholder {
    color: #b0b7c3 !important;
}
/* Ensure the baseweb container itself doesn't show a competing background */
.stTextArea [data-baseweb="textarea"],
.stTextArea > div > div {
    background: transparent !important;
    border: none !important;
}

/* ── Selectbox ── */
.stSelectbox [data-baseweb="select"] > div:first-child {
    border-radius: 8px !important;
    border: 1.5px solid #c9cfd8 !important;
    background: #ffffff !important;
    font-size: 0.875rem !important;
    color: #1f2937 !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stSelectbox [data-baseweb="select"] > div:first-child:focus-within {
    border-color: #374151 !important;
    box-shadow: 0 0 0 3px rgba(55,65,81,0.12) !important;
}
/* Text & icon colours inside the closed select box only (not the portal) */
.stSelectbox [data-baseweb="select"] > div:first-child span,
.stSelectbox [data-baseweb="select"] > div:first-child > div { color: #1f2937 !important; }

/* ── Multiselect ── */
.stMultiSelect [data-baseweb="select"] > div:first-child {
    border-radius: 8px !important;
    border: 1.5px solid #c9cfd8 !important;
    background: #ffffff !important;
    font-size: 0.875rem !important;
    color: #1f2937 !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stMultiSelect [data-baseweb="select"] > div:first-child:focus-within {
    border-color: #374151 !important;
    box-shadow: 0 0 0 3px rgba(55,65,81,0.12) !important;
}
.stMultiSelect input { color: #1f2937 !important; }
.stMultiSelect [data-baseweb="select"] span { color: #1f2937 !important; }

/* ── Dropdown portal (renders at body root, outside stApp) ── */
/* Option list container */
[data-baseweb="popover"] [data-baseweb="menu"],
[data-baseweb="select__dropdown"] {
    background: #ffffff !important;
    border: 1px solid #d1d5db !important;
    border-radius: 8px !important;
    box-shadow: 0 8px 24px rgba(0,0,0,0.12) !important;
}
/* Individual options */
[data-baseweb="menu"] [role="option"],
[data-baseweb="popover"] li {
    font-size: 0.875rem !important;
    color: #1f2937 !important;
    background: #ffffff !important;
    font-family: 'Inter', sans-serif !important;
}
/* Hovered option */
[data-baseweb="menu"] [role="option"]:hover,
[data-baseweb="popover"] li:hover {
    background: #f1f5f9 !important;
    color: #0f172a !important;
}
/* Selected/active option */
[data-baseweb="menu"] [aria-selected="true"],
[data-baseweb="menu"] [role="option"][data-highlighted] {
    background: #1f2937 !important;
    color: #ffffff !important;
}

/* ── Primary button ── */
div[data-testid="stMainBlockContainer"] .stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #1f2937 0%, #374151 100%) !important;
    color: #ffffff !important;
    border: none !important;
    padding: 12px 32px !important;
    border-radius: 9px !important;
    font-size: 0.9rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.02em !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 2px 8px rgba(31,41,55,0.22) !important;
}
div[data-testid="stMainBlockContainer"] .stButton > button[kind="primary"]:hover {
    background: linear-gradient(135deg, #111827 0%, #1f2937 100%) !important;
    color: #ffffff !important;
    box-shadow: 0 8px 24px rgba(0,0,0,0.2) !important;
    transform: translateY(-1px) !important;
}
div[data-testid="stMainBlockContainer"] .stButton > button[kind="primary"]:focus-visible {
    outline: 2px solid #374151 !important;
    outline-offset: 3px !important;
    box-shadow: 0 0 0 4px rgba(55,65,81,0.2) !important;
}
/* White text for primary buttons only — scoped to [kind="primary"] so
   secondary buttons (light bg) are NOT affected. */
div[data-testid="stMainBlockContainer"] .stButton > button[kind="primary"] p,
div[data-testid="stMainBlockContainer"] .stButton > button[kind="primary"] span {
    color: #ffffff !important;
}

/* ── Secondary button (default st.button with no type=) ── */
div[data-testid="stMainBlockContainer"] .stButton > button[kind="secondary"] {
    background: #ffffff !important;
    color: #1f2937 !important;
    border: 1.5px solid #c9cfd8 !important;
    border-radius: 9px !important;
    font-size: 0.875rem !important;
    font-weight: 500 !important;
    padding: 8px 18px !important;
    transition: all 0.2s ease !important;
    box-shadow: none !important;
}
div[data-testid="stMainBlockContainer"] .stButton > button[kind="secondary"]:hover {
    background: #f1f5f9 !important;
    border-color: #9ca3af !important;
    color: #0f172a !important;
    box-shadow: 0 2px 6px rgba(0,0,0,0.08) !important;
}
div[data-testid="stMainBlockContainer"] .stButton > button[kind="secondary"]:active {
    background: #e8ecf0 !important;
    border-color: #9ca3af !important;
    color: #0f172a !important;
    box-shadow: none !important;
    transform: translateY(1px) !important;
}
div[data-testid="stMainBlockContainer"] .stButton > button[kind="secondary"]:focus-visible {
    outline: 2px solid #374151 !important;
    outline-offset: 3px !important;
}
div[data-testid="stMainBlockContainer"] .stButton > button[kind="secondary"] p,
div[data-testid="stMainBlockContainer"] .stButton > button[kind="secondary"] span {
    color: #1f2937 !important;
}

/* ── Download button ── */
.stDownloadButton > button {
    background: linear-gradient(135deg, #1f2937 0%, #374151 100%) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 9px !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    padding: 12px 20px !important;
    width: 100% !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 2px 8px rgba(31,41,55,0.22) !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg, #111827 0%, #1f2937 100%) !important;
    color: #ffffff !important;
    box-shadow: 0 8px 24px rgba(0,0,0,0.2) !important;
    transform: translateY(-1px) !important;
}
.stDownloadButton > button:focus-visible {
    outline: 2px solid #374151 !important;
    outline-offset: 3px !important;
    box-shadow: 0 0 0 4px rgba(55,65,81,0.2) !important;
}
/* Same fix for download button inner text */
.stDownloadButton > button p,
.stDownloadButton > button span {
    color: #ffffff !important;
}

/* ── Expander (main area) ── */
details > summary > div > p {
    font-weight: 600 !important;
    color: #1f2937 !important;
    font-size: 0.875rem !important;
}
/* Chevron icon beside the summary label */
details > summary svg {
    color: #4b5563 !important;
    fill: #4b5563 !important;
}
details {
    border: 1px solid #d1d5db !important;
    border-radius: 10px !important;
    background: white !important;
    margin-bottom: 10px !important;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04) !important;
    transition: box-shadow 0.2s !important;
}
details[open] {
    box-shadow: 0 4px 12px rgba(0,0,0,0.08) !important;
}

/* ── Radio — rectangular pill buttons ── */
.stRadio > label {
    font-size: 0.875rem !important;
    font-weight: 600 !important;
    color: #374151 !important;
    letter-spacing: 0.01em !important;
}
.stRadio > div,
.stRadio [role="radiogroup"] {
    display: flex !important;
    flex-direction: row !important;
    align-items: center !important;
    flex-wrap: wrap !important;
    gap: 8px !important;
}
.stRadio > div > label,
.stRadio [role="radiogroup"] label {
    display: inline-flex !important;
    align-items: center !important;
    justify-content: center !important;
    padding: 6px 14px !important;
    height: 32px !important;
    border-radius: 6px !important;
    border: 1.5px solid #c9cfd8 !important;
    background: #ffffff !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    color: #374151 !important;
    cursor: pointer !important;
    transition: background 0.15s ease, border-color 0.15s ease, color 0.15s ease !important;
    margin: 0 !important;
    line-height: 1 !important;
    box-sizing: border-box !important;
}
.stRadio > div > label:hover,
.stRadio [role="radiogroup"] label:hover {
    background: #f1f5f9 !important;
    border-color: #9ca3af !important;
}
.stRadio > div > label:has(input:checked),
.stRadio [role="radiogroup"] label:has(input:checked) {
    background: #1f2937 !important;
    border-color: #1f2937 !important;
    color: #ffffff !important;
}
/* Ensure text and nested elements inherit white color on selected state */
.stRadio > div > label:has(input:checked) p,
.stRadio > div > label:has(input:checked) span,
.stRadio > div > label:has(input:checked) div,
.stRadio [role="radiogroup"] label:has(input:checked) p,
.stRadio [role="radiogroup"] label:has(input:checked) span,
.stRadio [role="radiogroup"] label:has(input:checked) div {
    color: #ffffff !important;
}
.stRadio label p {
    margin: 0 !important;
    line-height: 1 !important;
    color: inherit !important;
}
/* Radio keyboard focus */
.stRadio > div > label:has(input:focus-visible),
.stRadio [role="radiogroup"] label:has(input:focus-visible) {
    outline: 2px solid #374151 !important;
    outline-offset: 2px !important;
}

/* ── Date input ── */
.stDateInput input {
    border-radius: 8px !important;
    border: 1.5px solid #c9cfd8 !important;
    background: #ffffff !important;
    font-size: 0.875rem !important;
    color: #1f2937 !important;
    padding: 8px 12px !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stDateInput input:focus {
    border-color: #374151 !important;
    box-shadow: 0 0 0 3px rgba(55,65,81,0.12) !important;
    outline: none !important;
}
.stDateInput input::placeholder { color: #b0b7c3 !important; }
.stDateInput [data-baseweb="input"],
.stDateInput > div > div {
    background: transparent !important;
    border: none !important;
}

/* ── Labels ── */
.stTextInput label, .stTextArea label, .stSelectbox label,
.stMultiSelect label, .stNumberInput label, .stDateInput label,
.stRadio > label {
    font-size: 0.875rem !important;
    font-weight: 600 !important;
    color: #374151 !important;
    letter-spacing: 0.01em !important;
}

/* ── Dividers ── */
hr {
    border: none !important;
    border-top: 1px solid #d1d5db !important;
    margin: 18px 0 !important;
}

/* ── Spinner ── */
.stSpinner > div { border-top-color: #374151 !important; }

/* ── Alerts (info / success / warning / error) ── */
.stAlert {
    border-radius: 10px !important;
    font-size: 0.875rem !important;
    font-family: 'Inter', sans-serif !important;
}
/* Let each alert type keep its own colour — the global "> div p" rule would
   otherwise force all alert text to #1f2937, killing green/red/yellow coding. */
.stAlert p, .stAlert span, .stAlert a {
    color: inherit !important;
    font-size: 0.875rem !important;
}
[data-testid="stNotification"] {
    border-radius: 10px !important;
}

/* ── Markdown bold ── */
.stMarkdown strong { color: #111827 !important; font-weight: 600 !important; }

/* ── Tag chips in multiselect ── */
.stMultiSelect span[data-baseweb="tag"] {
    background-color: #f1f5f9 !important;
    border-radius: 6px !important;
    color: #1f2937 !important;
    font-size: 0.78rem !important;
}
.stMultiSelect span[data-baseweb="tag"] button { color: #6b7280 !important; }
.stMultiSelect input::placeholder { color: #b0b7c3 !important; }

/* ── Number input — full styling including stepper buttons ── */
.stNumberInput input {
    background: #ffffff !important;
    border: 1.5px solid #c9cfd8 !important;
    border-radius: 8px !important;
    color: #1f2937 !important;
    font-size: 0.875rem !important;
    padding: 8px 12px !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stNumberInput input:focus {
    border-color: #374151 !important;
    box-shadow: 0 0 0 3px rgba(55,65,81,0.12) !important;
    outline: none !important;
}
.stNumberInput input::placeholder { color: #b0b7c3 !important; }
/* Stepper +/- buttons */
.stNumberInput [data-baseweb="input"] button {
    background: #f8fafc !important;
    border-left: 1px solid #c9cfd8 !important;
    color: #374151 !important;
}
.stNumberInput [data-baseweb="input"] button:hover {
    background: #f1f5f9 !important;
    color: #1f2937 !important;
}
/* Remove competing wrapper border */
.stNumberInput [data-baseweb="input"] {
    border: none !important;
    background: transparent !important;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────

if "page"             not in st.session_state: st.session_state["page"]             = "main"
if "authenticated"    not in st.session_state: st.session_state["authenticated"]    = False
if "current_user"     not in st.session_state: st.session_state["current_user"]     = None
if "display_name"     not in st.session_state: st.session_state["display_name"]     = None
if "failed_attempts"        not in st.session_state: st.session_state["failed_attempts"]        = 0
if "lockout_until"          not in st.session_state: st.session_state["lockout_until"]          = 0.0
if "login_username_preview" not in st.session_state: st.session_state["login_username_preview"] = ""
if "cc_show_results"        not in st.session_state: st.session_state["cc_show_results"]        = False
if "analytics_sid"          not in st.session_state: st.session_state["analytics_sid"]          = str(uuid.uuid4())
if "analytics_last_page"    not in st.session_state: st.session_state["analytics_last_page"]    = ""
if "ls_write_token"         not in st.session_state: st.session_state["ls_write_token"]         = ""
if "ls_clear"               not in st.session_state: st.session_state["ls_clear"]               = False
if "_editing_submission_file" not in st.session_state: st.session_state["_editing_submission_file"] = None

# Session restoration from localStorage is handled by _session_mgr component below.


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

@st.cache_data
def get_base64_image(path):
    """Read and base64-encode an image file. Result is cached across rerenders."""
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()


def celebrate(message="Done!", sub=""):
    """Balloons + confetti burst + animated success banner."""
    st.balloons()
    components.html("""
    <script src="https://cdn.jsdelivr.net/npm/canvas-confetti@1.9.3/dist/confetti.browser.min.js"></script>
    <script>
    (function(){
        var end = Date.now() + 2200;
        var colors = ['#1f2937','#374151','#6b7280','#9ca3af','#d1d5db'];
        (function frame(){
            confetti({ particleCount:6, angle:60,  spread:55, origin:{x:0}, colors:colors });
            confetti({ particleCount:6, angle:120, spread:55, origin:{x:1}, colors:colors });
            if (Date.now() < end) requestAnimationFrame(frame);
        }());
    })();
    </script>
    """, height=0)
    st.markdown(f"""
    <style>
    @keyframes slideDown {{
        from {{ opacity: 0; transform: translateY(-16px); }}
        to   {{ opacity: 1; transform: translateY(0); }}
    }}
    </style>
    <div style="
        animation: slideDown 0.4s cubic-bezier(0.34,1.56,0.64,1) forwards;
        background: linear-gradient(135deg, #f0fdf4 0%, #dcfce7 100%);
        border: 1px solid #86efac;
        border-left: 5px solid #16a34a;
        border-radius: 12px;
        padding: 15px 20px;
        margin: 14px 0 10px 0;
        display: flex; align-items: center; gap: 14px;
        box-shadow: 0 4px 16px rgba(22,163,74,0.12);
        font-family: 'Inter', sans-serif;
    ">
        <div style="font-size:1.7rem; flex-shrink:0;">✅</div>
        <div>
            <div style="font-weight:700;color:#15803d;font-size:0.95rem;letter-spacing:0.01em;">{_h(message)}</div>
            {"" if not sub else f'<div style="color:#166534;font-size:0.8rem;margin-top:4px;opacity:0.85;">{_h(sub)}</div>'}
        </div>
    </div>
    """, unsafe_allow_html=True)


def section_header(icon, title):
    """Styled section header bar used inside form pages."""
    st.markdown(f"""
    <div style="
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
        border-left: 4px solid #1f2937;
        border-radius: 0 10px 10px 0;
        color: #0f172a;
        padding: 11px 18px;
        font-weight: 700;
        font-size: 0.875rem;
        margin: 30px 0 16px 0;
        letter-spacing: 0.02em;
        display: flex;
        align-items: center;
        gap: 9px;
        box-shadow: 0 1px 4px rgba(0,0,0,0.06);
        font-family: 'Inter', sans-serif;
    ">{icon}&ensp;{title}</div>
    """, unsafe_allow_html=True)


def page_title(title, subtitle=""):
    st.markdown(f"""
    <div style="padding: 4px 0 22px 0; margin-bottom: 4px;">
        <div style="font-size: 1.65rem; font-weight: 700; color: #0f172a; line-height: 1.2; letter-spacing: -0.025em; font-family: 'Inter', sans-serif;">{_h(title)}</div>
        {"" if not subtitle else f'<div style="font-size:0.875rem;color:#6b7280;margin-top:7px;font-weight:400;line-height:1.5;">{_h(subtitle)}</div>'}
        <div style="height:3px;background:linear-gradient(90deg,#1f2937 0%,#6b7280 55%,transparent 100%);border-radius:2px;margin-top:16px;"></div>
    </div>
    """, unsafe_allow_html=True)


def info_row(label, value):
    # _h() escapes user-supplied value to prevent XSS via unsafe_allow_html
    st.markdown(f"""
    <div style="padding: 8px 10px 7px 10px; border-bottom: 1px solid #e5e7eb;">
        <div style="color:#6b7280; font-size:0.66rem; text-transform:uppercase; letter-spacing:0.08em; font-weight:700; margin-bottom:3px; font-family:'Inter',sans-serif;">{_h(label)}</div>
        <div style="color:#111827; font-size:0.875rem; font-weight:500; font-family:'Inter',sans-serif;">{_h(value)}</div>
    </div>
    """, unsafe_allow_html=True)


def frequency_selector(label, key_prefix):
    col1, col2 = st.columns([1, 1])
    with col1:
        freq = st.selectbox(f"{label} — Frequency", ["Daily", "Weekly", "Monthly", "Hourly"], key=f"{key_prefix}_freq")
    hourly_count = None
    daily_count = None
    if freq == "Hourly":
        with col2:
            hourly_count = st.number_input("Times / day", min_value=1, key=f"{key_prefix}_hourly")
    elif freq == "Daily":
        with col2:
            daily_count = st.radio("Times / day", [1, 2], key=f"{key_prefix}_daily", horizontal=True)
    return freq, hourly_count, daily_count


def calculate_risk(freq_string, volume_string):
    if not freq_string or not volume_string:
        return None
    freq_score = 1
    if "Hourly" in freq_string:
        try:
            times = int(freq_string.split("(")[1].split()[0])
            freq_score = 3 if times > 6 else 2
        except (IndexError, ValueError):
            freq_score = 2
    vol_score = 1
    try:
        volume = int(str(volume_string).replace(",", ""))
        vol_score = 1 if volume <= 10_000 else (2 if volume <= 50_000 else 3)
    except ValueError:
        pass
    total = freq_score + vol_score
    return "LOW" if total <= 2 else ("MODERATE" if total <= 4 else "CRITICAL")


PREDEFINED_DOMAINS = ["swiggy.com", "blinkit.com", "zeptonow.com", "amazon.in", "flipkart.com"]


def _section_label(text: str) -> None:
    """Render a bold section label as HTML.

    Use this instead of st.markdown(f"**{text}**") whenever `text` is a
    variable or contains an asterisk — trailing/embedded asterisks break
    Markdown bold syntax (e.g. "**Domains ***" is invalid).
    """
    if text.endswith(" *"):
        base = text[:-2]
        html = (
            f'<span style="font-weight:600;color:#374151;font-size:0.875rem;">'
            f'{base} <span style="color:#dc2626;font-weight:700;">*</span></span>'
        )
    else:
        html = f'<span style="font-weight:600;color:#374151;font-size:0.875rem;">{text}</span>'
    st.markdown(html, unsafe_allow_html=True)


def domain_selector(label, key_prefix):
    _section_label(label)
    col1, col2 = st.columns([3, 1])
    with col1:
        selected = st.multiselect(label, PREDEFINED_DOMAINS, key=f"{key_prefix}_domains", label_visibility="collapsed")
    with col2:
        custom = st.text_input("Custom", placeholder="+ Add domain", key=f"{key_prefix}_custom_domain", label_visibility="collapsed")
    domains = selected + ([custom.strip()] if custom.strip() else [])
    return ", ".join(domains) if domains else "", domains


def _safe_key(s):
    """Return a Streamlit-safe widget key fragment from an arbitrary string."""
    return re.sub(r'[^a-zA-Z0-9]', '_', s)


# ── Per-module crawl-config helpers ──────────────────────────────────────────
# Each helper renders config widgets and returns a flat dict.
# key_suffix is appended to every widget key so per-domain instances don't clash.

def _pt_crawl_config(key_suffix=""):
    cfg = {}
    crawl_type = st.radio(
        "Crawl Type", ["Category-based (Category_ES)", "Input-based (URL/Input driven)", "Products Only"],
        horizontal=True, key=f"pt_crawl_type{key_suffix}"
    )
    cfg["Crawl Type"] = crawl_type

    st.markdown("**Overall Crawl Frequency**")
    freq, hourly, daily = frequency_selector("Overall", f"pt_overall{key_suffix}")
    cfg["Overall Frequency"] = f"{freq} ({hourly} times/day)" if hourly else (f"{freq} ({daily}x/day)" if daily and daily > 1 else freq)

    if crawl_type == "Products Only":
        st.markdown("---")
        st.markdown("##### C) Products Only Configuration")

        st.markdown("**Products Crawl Frequency**")
        pf, ph, pd = frequency_selector("Products Crawl", f"pt_prodonly{key_suffix}")
        cfg["Products Crawl Frequency"] = f"{pf} ({ph} times/day)" if ph else (f"{pf} ({pd}x/day)" if pd and pd > 1 else pf)
        if ph:
            cfg["Hourly Crawl Timings"] = st.text_input(
                "Specify crawl hours", placeholder="e.g., 9 AM, 12 PM, 3 PM, 6 PM",
                key=f"pt_prodonly_hourly_timings{key_suffix}"
            )

        st.markdown("**Inputs**")
        cfg["Sample Input URLs"] = st.text_area("Sample Product URLs", placeholder="If client inputs not available, provide testing URLs", key=f"pt_prodonly_sample_urls{key_suffix}")
        inp_status = st.radio("Client Inputs Status", ["Not Yet Provided", "Available — See Sheet Link Below"], key=f"pt_prodonly_inputs_status{key_suffix}", horizontal=True)
        if inp_status == "Not Yet Provided":
            cfg["Client Inputs Expected Date"] = str(st.date_input("Expected delivery date for inputs", key=f"pt_prodonly_inputs_expected_date{key_suffix}"))
        else:
            cfg["Client Inputs Sheet Link"] = st.text_input("Sheet Link with client inputs", key=f"pt_prodonly_inputs_sheet_link{key_suffix}")

        st.markdown("**Location Dependency**")
        is_pincode = st.radio("Pincode / Zipcode based?", ["Yes", "No"], key=f"pt_prodonly_pincode_based{key_suffix}", horizontal=True)
        cfg["Pincode Based"] = is_pincode
        if is_pincode == "Yes":
            c1, c2 = st.columns(2)
            with c1:
                cfg["Sample Pincode"] = st.text_input("Sample Pincode", placeholder="e.g., 110001, 560001", key=f"pt_prodonly_sample_pincode{key_suffix}")
            with c2:
                cfg["Client Pincode List Link"] = st.text_input("Pincode list link (if available)", key=f"pt_prodonly_pincode_list_link{key_suffix}")

        st.markdown("**Volume & Output**")
        c1, c2 = st.columns(2)
        with c1:
            cfg["Expected Volume"] = st.text_input("Expected Volume / day", placeholder="e.g., 50,000 products", key=f"pt_prodonly_expected_volume{key_suffix}")
        with c2:
            cfg["Screenshot Required"] = st.radio("Screenshot Required?", ["Yes", "No"], key=f"pt_prodonly_screenshot{key_suffix}", horizontal=True)

    elif crawl_type == "Category-based (Category_ES)":
        st.markdown("---")
        st.markdown("##### A) Category_ES Configuration")

        st.markdown("**Index Frequency**")
        c1, c2 = st.columns(2)
        with c1:
            pf, ph, pd = frequency_selector("Products Index", f"pt_prod{key_suffix}")
            cfg["Products Index Frequency"] = f"{pf} ({ph} times/day)" if ph else (f"{pf} ({pd}x/day)" if pd and pd > 1 else pf)
        with c2:
            tf, th, td = frequency_selector("Trends Index", f"pt_trend{key_suffix}")
            cfg["Trends Index Frequency"] = f"{tf} ({th} times/day)" if th else (f"{tf} ({td}x/day)" if td and td > 1 else tf)

        if ph or th:
            cfg["Hourly Crawl Timings"] = st.text_input(
                "Specify crawl hours", placeholder="e.g., 9 AM, 12 PM, 3 PM, 6 PM",
                key=f"pt_hourly_timings{key_suffix}"
            )

        st.markdown("**Trends Configuration**")
        c1, c2 = st.columns(2)
        with c1:
            cfg["No of RSS Crawls"] = st.number_input("Number of RSS crawls into Trends", min_value=0, key=f"pt_rss_crawls{key_suffix}")
        with c2:
            cfg["Expected Data Push Volume"] = st.text_input("Products volume to push into Trends", key=f"pt_data_push_volume{key_suffix}")

        st.markdown("**Category Details**")
        cfg["Sample Category List"] = st.text_area(
            "Sample Category List", placeholder="e.g., Electronics, Fashion, Home & Kitchen",
            key=f"pt_sample_category_list{key_suffix}"
        )
        cat_status = st.radio("Is final category list available?", ["Yes", "No"], key=f"pt_category_status{key_suffix}", horizontal=True)
        if cat_status == "Yes":
            cfg["Client Category Sheet Link"] = st.text_input("Category Sheet Link", key=f"pt_category_sheet_link{key_suffix}")
        else:
            cfg["Client Category Expected Date"] = str(st.date_input("Expected date for category list", key=f"pt_category_expected_date{key_suffix}"))

    elif crawl_type == "Input-based (URL/Input driven)":
        st.markdown("---")
        st.markdown("##### B) Input-Based Configuration")

        st.markdown("**Products Crawl**")
        need_product = st.radio("Products crawl required?", ["Yes", "No"], key=f"pt_input_products_needed{key_suffix}", horizontal=True)
        cfg["Products Crawl Needed"] = need_product
        if need_product == "Yes":
            pf, ph, pd = frequency_selector("Products Crawl", f"pt_input_prod{key_suffix}")
            cfg["Products Crawl Frequency"] = f"{pf} ({ph} times/day)" if ph else (f"{pf} ({pd}x/day)" if pd and pd > 1 else pf)

        st.markdown("**Trends Crawl**")
        tf, th, td = frequency_selector("Trends Crawl", f"pt_input_trend{key_suffix}")
        cfg["Trends Crawl Frequency"] = f"{tf} ({th} times/day)" if th else (f"{tf} ({td}x/day)" if td and td > 1 else tf)
        if th:
            cfg["Trends Hourly Timings"] = st.text_input(
                "Specify timing if hourly", placeholder="e.g., 10 AM, 2 PM, 6 PM, 10 PM",
                key=f"pt_trends_hourly_timings{key_suffix}"
            )

        st.markdown("**Inputs**")
        cfg["Sample Input URLs"] = st.text_area("Sample Input URLs", placeholder="If client inputs not available, provide testing URLs", key=f"pt_sample_input_urls{key_suffix}")
        inp_status = st.radio("Client Inputs Status", ["Not Yet Provided", "Available — See Sheet Link Below"], key=f"pt_inputs_status{key_suffix}", horizontal=True)
        if inp_status == "Not Yet Provided":
            cfg["Client Inputs Expected Date"] = str(st.date_input("Expected delivery date for inputs", key=f"pt_inputs_expected_date{key_suffix}"))
        else:
            cfg["Client Inputs Sheet Link"] = st.text_input("Sheet Link with client inputs", key=f"pt_inputs_sheet_link{key_suffix}")

        st.markdown("**Location Dependency**")
        is_pincode = st.radio("Pincode / Zipcode based?", ["Yes", "No"], key=f"pt_pincode_based{key_suffix}", horizontal=True)
        cfg["Pincode Based"] = is_pincode
        if is_pincode == "Yes":
            c1, c2 = st.columns(2)
            with c1:
                cfg["Sample Pincode"] = st.text_input("Sample Pincode", placeholder="e.g., 110001, 560001", key=f"pt_sample_pincode{key_suffix}")
            with c2:
                cfg["Client Pincode List Link"] = st.text_input("Pincode list link (if available)", key=f"pt_pincode_list_link{key_suffix}")

        st.markdown("**Crawl Duration**")
        c1, c2 = st.columns(2)
        with c1:
            cfg["Crawl Start Date"] = str(st.date_input("Start Date", key=f"pt_crawl_start_date{key_suffix}"))
        with c2:
            cfg["Crawl End Date"] = str(st.date_input("End Date", key=f"pt_crawl_end_date{key_suffix}"))

        st.markdown("**Volume & Output**")
        c1, c2 = st.columns(2)
        with c1:
            cfg["Expected Volume"] = st.text_input("Expected Volume / day", placeholder="e.g., 50,000 products", key=f"pt_expected_volume{key_suffix}")
        with c2:
            cfg["Screenshot Required"] = st.radio("Screenshot Required?", ["Yes", "No"], key=f"pt_screenshot{key_suffix}", horizontal=True)

    return cfg


def _sos_crawl_config(key_suffix=""):
    cfg = {}
    cfg["Zipcode Required"] = st.radio("Zipcode required?", ["Yes", "No"], horizontal=True, key=f"sos_zipcode_required{key_suffix}")
    if cfg["Zipcode Required"] == "Yes":
        cfg["Pincode List"] = st.text_area(
            "Pincode list (comma-separated or sheet link)",
            placeholder="e.g., 110001, 560001, 400001",
            key=f"sos_pincode_list{key_suffix}"
        )
        st.caption("Format: comma-separated pincodes (e.g., 110001, 560001) or paste a sheet link.")

    st.markdown("**Crawl Depth**")
    c1, c2 = st.columns(2)
    with c1:
        cfg["No. of Pages per Keyword"] = st.number_input("Pages per keyword", min_value=1, value=1, key=f"sos_pages{key_suffix}")
    with c2:
        cfg["No. of Products per Keyword"] = st.number_input("Products per keyword", min_value=1, value=10, key=f"sos_products{key_suffix}")

    st.markdown("**Crawl Frequency**")
    freq, hourly, daily = frequency_selector("SOS Crawl", f"sos{key_suffix}")
    cfg["Frequency"] = f"{freq} ({hourly} times/day)" if hourly else (f"{freq} ({daily}x/day)" if daily and daily > 1 else freq)
    return cfg


def _rev_crawl_config(key_suffix=""):
    cfg = {}
    st.markdown("**Review Source Type**")
    cfg["Input Sources"] = st.multiselect(
        "Where to pull review inputs from *",
        ["From Products Index", "From Trends Index", "From Review Input URLs", "Category-based Reviews Crawl"],
        key=f"rev_source{key_suffix}",
    )
    if "From Review Input URLs" in cfg["Input Sources"]:
        cfg["Sample Review URLs"] = st.text_area("Sample review page URLs", placeholder="Provide product review page URLs", key=f"rev_sample_urls{key_suffix}")

    st.markdown("**Frequency**")
    c1, c2 = st.columns(2)
    with c1:
        freq, hourly, daily = frequency_selector("Reviews Crawl", f"rev{key_suffix}")
        cfg["Frequency"] = f"{freq} ({hourly} times/day)" if hourly else (f"{freq} ({daily}x/day)" if daily and daily > 1 else freq)
    if hourly:
        with c2:
            cfg["Hourly Timings"] = st.text_input("Timing if hourly", placeholder="e.g., 8 AM, 12 PM, 6 PM, 10 PM", key=f"rev_hourly_timings{key_suffix}")
    return cfg


def _pv_crawl_config(key_suffix=""):
    cfg = {}
    st.markdown("**Frequency**")
    freq, hourly, daily = frequency_selector("Price Violation Crawl", f"pv{key_suffix}")
    cfg["Frequency"] = f"{freq} ({hourly} times/day)" if hourly else (f"{freq} ({daily}x/day)" if daily and daily > 1 else freq)

    st.markdown("**Inputs**")
    cfg["Product URL List"] = st.text_area("Product URL list to monitor *", placeholder="Sample product URLs", key=f"pv_product_url_list{key_suffix}")

    cfg["Zipcode Required"] = st.radio("Zipcode required?", ["Yes", "No"], horizontal=True, key=f"pv_zipcode_required{key_suffix}")
    if cfg["Zipcode Required"] == "Yes":
        cfg["Zipcode List"] = st.text_area("Zipcode list", placeholder="e.g., 110001, 560001, 400001", key=f"pv_zipcode_list{key_suffix}")

    cfg["Price Violation Condition"] = st.text_area(
        "Violation condition / rule",
        placeholder="e.g., MRP > X, Discount < Y%, price diff > 15%",
        key=f"pv_violation_condition{key_suffix}"
    )
    c1, c2 = st.columns(2)
    with c1:
        cfg["Sample Inputs Sheet Link"] = st.text_input("Sample inputs sheet link", placeholder="Link to sample data", key=f"pv_sample_inputs_link{key_suffix}")
    with c2:
        cfg["Screenshot Required"] = st.radio("Screenshot Required?", ["Yes", "No"], key=f"pv_screenshot{key_suffix}", horizontal=True)
    return cfg


def _storeid_crawl_config(key_suffix=""):
    cfg = {}
    c1, _ = st.columns(2)
    with c1:
        cfg["Specific Location Required"] = st.radio(
            "Specific store locations needed?", ["No", "Yes"], horizontal=True, key=f"storeid_location{key_suffix}"
        )
    if cfg["Specific Location Required"] == "Yes":
        cfg["Location Details"] = st.text_area("Location details", placeholder="e.g., Bangalore, Mumbai, Delhi", key=f"storeid_location_details{key_suffix}")

    storeid_status = st.radio("Specific Pincode list available?", ["Yes", "No"], horizontal=True, key=f"storeid_list_status{key_suffix}")
    if storeid_status == "Yes":
        cfg["Specific Pincode List Link"] = st.text_input("Pincode list link", key=f"storeid_pincode_list_link{key_suffix}")
    return cfg


def _festive_schedule_config(key_suffix=""):
    cfg = {}
    c1, c2, c3 = st.columns(3)
    with c1:
        cfg["Frequency Per Day"] = st.number_input("Frequency / day", min_value=1, value=1, key=f"festive_freq{key_suffix}")
    with c2:
        cfg["Start Date"] = str(st.date_input("Start Date", key=f"festive_start{key_suffix}"))
    with c3:
        cfg["End Date"] = str(st.date_input("End Date", key=f"festive_end{key_suffix}"))
    return cfg


def _apply_domain_config(base_dict, config_mode_key, domain_list, config_fn):
    """
    If multiple domains: ask same/different. Render config once (common) or per domain.
    Merges results into base_dict (per-domain keys prefixed with domain name).
    """
    if len(domain_list) > 1:
        config_mode = st.radio(
            "Crawl configuration",
            ["Same config for all domains", "Different config per domain"],
            key=f"{config_mode_key}_mode",
            horizontal=True,
        )
    else:
        config_mode = "Same config for all domains"

    if config_mode == "Same config for all domains":
        cfg = config_fn("")
        base_dict.update(cfg)
    else:
        for domain in domain_list:
            sk = _safe_key(domain)
            with st.expander(f"Config for: **{domain}**", expanded=True):
                cfg = config_fn(f"_{sk}")
            for k, v in cfg.items():
                base_dict[f"{domain} — {k}"] = v


def validate_required(client_name):
    if not client_name:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #fffbeb 0%, #fef9ec 100%);
            border: 1px solid #fcd34d;
            border-left: 4px solid #f59e0b;
            border-radius: 10px;
            padding: 13px 18px;
            color: #78350f;
            font-size: 0.875rem;
            margin: 6px 0 22px 0;
            display: flex;
            align-items: center;
            gap: 10px;
            box-shadow: 0 1px 4px rgba(245,158,11,0.12);
            font-family: 'Inter', sans-serif;
        ">
            <span style="font-size:1.1rem;">⚠️</span>
            <span>Please enter a <strong>Client Name</strong> above to unlock the full form.</span>
        </div>
        """, unsafe_allow_html=True)
        st.stop()


def render_summary(data):
    st.markdown("""
    <div style="
        background: linear-gradient(135deg, #1f2937 0%, #374151 100%);
        border-radius: 10px;
        padding: 14px 18px;
        margin-bottom: 18px;
        box-shadow: 0 4px 12px rgba(31,41,55,0.2);
    ">
        <div style="font-size: 0.9rem; font-weight: 700; color: #ffffff; letter-spacing: 0.02em; font-family:'Inter',sans-serif;">📋 Live Summary</div>
        <div style="font-size: 0.73rem; color: #9ca3af; margin-top: 3px; font-family:'Inter',sans-serif;">Auto-updates as you fill the form</div>
    </div>
    """, unsafe_allow_html=True)

    has_content = any(
        any(v not in ["", None, [], {}] for v in c.values())
        for c in data.values()
    )

    if not has_content:
        st.markdown("""
        <div style="
            text-align:center; padding:44px 20px;
            background:white; border-radius:12px; border:2px dashed #d1d5db;
            box-shadow: 0 1px 4px rgba(0,0,0,0.04);
        ">
            <div style="font-size:2.4rem; margin-bottom:12px; opacity:0.55;">📝</div>
            <div style="font-size:0.875rem; font-weight:600; color:#64748b; margin-bottom:5px; font-family:'Inter',sans-serif;">Nothing here yet</div>
            <div style="font-size:0.78rem; color:#6b7280; line-height:1.6; font-family:'Inter',sans-serif;">
                Start filling the form<br>to preview a summary here
            </div>
        </div>
        """, unsafe_allow_html=True)
        return

    for section, content in data.items():
        filled = {k: v for k, v in content.items() if v not in ["", None, [], {}]}
        if not filled:
            continue
        with st.expander(f"**{section}**", expanded=True):
            for k, v in filled.items():
                info_row(k, v)

    if "Products + Trends" in data:
        pt = data["Products + Trends"]
        # Support both common config ("Overall Frequency") and per-domain config ("domain — Overall Frequency")
        _overall_freq = pt.get("Overall Frequency") or next(
            (v for k, v in pt.items() if k.endswith("Overall Frequency")), None
        )
        _expected_vol = pt.get("Expected Volume") or next(
            (v for k, v in pt.items() if k.endswith("Expected Volume")), None
        )
        risk = calculate_risk(_overall_freq, _expected_vol)
        if risk:
            st.markdown("---")
            st.markdown("**⚡ Crawl Load Risk**")
            if risk == "LOW":
                st.success("**LOW** — Infrastructure load is safe.")
            elif risk == "MODERATE":
                st.warning("**MODERATE** — Monitor scaling & proxy usage.")
            else:
                st.error("**CRITICAL** — High infra saturation risk.")


def generate_pdf(data, client_name):
    """Build a formatted PDF from form_data. All user values are XML-escaped
    before being passed to ReportLab Paragraph to prevent markup injection."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
    from reportlab.lib import pagesizes  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY  # type: ignore
    from reportlab.lib.colors import HexColor  # type: ignore
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=pagesizes.A4,
        topMargin=0.5*inch, bottomMargin=0.6*inch,
        leftMargin=0.6*inch, rightMargin=0.6*inch,
        title=f"{client_name} — Requirement Handling Form",
        author="42Signals",
    )
    styles = getSampleStyleSheet()
    el = []

    title_s = ParagraphStyle("T", parent=styles["Heading1"], fontSize=20, textColor=HexColor("#1f2937"),
                              alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=3)
    sub_s   = ParagraphStyle("S", parent=styles["Normal"], fontSize=9, textColor=HexColor("#6b7280"),
                              alignment=TA_CENTER, spaceAfter=16)
    sec_s   = ParagraphStyle("H", parent=styles["Heading2"], fontSize=10.5, textColor=HexColor("#1f2937"),
                              backColor=HexColor("#f3f4f6"), leftIndent=8, fontName="Helvetica-Bold",
                              spaceBefore=6, spaceAfter=4)
    key_s   = ParagraphStyle("K", parent=styles["Normal"], fontSize=8, textColor=HexColor("#6b7280"),
                              fontName="Helvetica-Bold")
    val_s   = ParagraphStyle("V", parent=styles["Normal"], fontSize=9, textColor=HexColor("#111827"),
                              wordWrap="CJK", alignment=TA_JUSTIFY)

    try:
        if os.path.exists(LOGO_PATH):
            logo = Image(LOGO_PATH, width=0.9*inch, height=0.72*inch)
            hdr = Table([[logo, Paragraph("<b>Requirement Handling Form</b>", title_s)]],
                        colWidths=[1.1*inch, 5.9*inch])
            hdr.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ]))
            el.append(hdr)
        else:
            el.append(Paragraph("<b>Requirement Handling Form</b>", title_s))
    except Exception:
        el.append(Paragraph("<b>Requirement Handling Form</b>", title_s))

    # Escape client_name before embedding in ReportLab XML
    safe_client = _html_mod.escape(str(client_name), quote=True)
    el.append(Paragraph(
        f"Client: {safe_client} &nbsp;|&nbsp; Generated: {date.today().strftime('%d %b %Y')}",
        sub_s
    ))
    el.append(Spacer(1, 0.1*inch))

    def _val_paragraph(v):
        if not v:
            return Paragraph("—", val_s)
        s = str(v)
        # detect URLs and render as clickable hyperlinks
        import re as _re
        url_pat = _re.compile(r'(https?://\S+)')
        parts = url_pat.split(s)
        if len(parts) == 1:
            return Paragraph(_html_mod.escape(s), val_s)
        xml = ""
        for part in parts:
            if url_pat.match(part):
                esc_url = _html_mod.escape(part, quote=True)
                xml += f'<a href="{esc_url}" color="#2563eb">{esc_url}</a>'
            else:
                xml += _html_mod.escape(part)
        return Paragraph(xml, val_s)

    row_colors = [HexColor("#f9fafb"), HexColor("#ffffff")]
    ci = 0
    for section, content in data.items():
        el.append(Paragraph(f"  {_html_mod.escape(str(section))}", sec_s))
        el.append(Spacer(1, 0.04*inch))
        rows = [
            [Paragraph(_html_mod.escape(str(k)), key_s),
             _val_paragraph(v)]
            for k, v in content.items()
        ]
        if rows:
            t = Table(rows, colWidths=[1.9*inch, 4.1*inch])
            t.setStyle(TableStyle([
                ("GRID",         (0, 0), (-1, -1), 0.4, HexColor("#e5e7eb")),
                ("VALIGN",       (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND",   (0, 0), (-1, -1), row_colors[ci % 2]),
                ("BACKGROUND",   (0, 0), (0, -1),  HexColor("#f3f4f6")),
                ("TOPPADDING",   (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 7),
                ("LEFTPADDING",  (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ]))
            el.append(t)
            el.append(Spacer(1, 0.14*inch))
            ci += 1

    doc.build(el)
    buffer.seek(0)
    return buffer


# ─────────────────────────────────────────────────────────────────────────────
# LOGIN PAGE
# ─────────────────────────────────────────────────────────────────────────────

def render_login():
    """Full-page sign-in screen — dark gradient background with white card."""

    # ── Page-level CSS (dark theme, animated orbs, card shell) ───────────
    st.markdown("""
    <style>
    @keyframes orbFloat {
        0%   { opacity:.5;  transform:scale(1)    translate(0,    0);    }
        100% { opacity:.85; transform:scale(1.08) translate(1.5%, 1.5%); }
    }
    @keyframes cardIn {
        from { opacity:0; transform:translateY(28px) scale(.97); }
        to   { opacity:1; transform:translateY(0)    scale(1);   }
    }
    @keyframes shimmer {
        0%   { background-position: -400px 0; }
        100% { background-position:  400px 0; }
    }

    /* ── Full-page dark background ── */
    .stApp {
        background: linear-gradient(145deg, #0a0f1e 0%, #0f172a 40%, #111827 100%) !important;
        min-height: 100vh;
        position: relative;
        overflow: hidden;
    }

    /* ── Animated gradient orbs ── */
    .stApp::before {
        content: '';
        position: fixed;
        inset: 0;
        background:
            radial-gradient(ellipse at 12% 22%,  rgba(99,102,241,.22) 0%, transparent 42%),
            radial-gradient(ellipse at 88% 78%,  rgba(16,185,129,.15) 0%, transparent 42%),
            radial-gradient(ellipse at 60% 5%,   rgba(59,130,246,.12) 0%, transparent 38%),
            radial-gradient(ellipse at 30% 90%,  rgba(139,92,246,.1)  0%, transparent 38%);
        animation: orbFloat 11s ease-in-out infinite alternate;
        pointer-events: none;
        z-index: 0;
    }

    /* ── Hide all chrome ── */
    section[data-testid="stSidebar"] { display:none !important; }
    header   { display:none !important; }
    footer   { display:none !important; }
    #MainMenu { display:none !important; }

    /* ── Card shell ── */
    .block-container {
        max-width: 450px !important;
        margin-top: 6vh !important;
        padding: 0 !important;
        background: #ffffff !important;
        border-radius: 24px !important;
        box-shadow:
            0 32px 80px rgba(0,0,0,.55),
            0  0   0  1px rgba(255,255,255,.07),
            inset 0 1px 0 rgba(255,255,255,.9) !important;
        position: relative !important;
        z-index: 1 !important;
        animation: cardIn .6s cubic-bezier(.34,1.56,.64,1) both !important;
    }

    /* ── Inner padding for Streamlit widgets (inputs, button, alerts) ── */
    [data-testid="stVerticalBlock"] {
        padding: 0 32px 24px 32px !important;
    }

    /* ── Login inputs (broad selectors to cover Streamlit's nested wrappers) ── */
    .stTextInput input,
    .stTextInput > div > div > input {
        background: #f8fafc !important;
        border: 1.5px solid #cbd5e1 !important;
        border-radius: 10px !important;
        font-size: 0.9rem !important;
        color: #0f172a !important;
        padding: 10px 14px !important;
        transition: all .2s !important;
    }
    .stTextInput input:focus,
    .stTextInput > div > div > input:focus {
        background: #ffffff !important;
        border-color: #374151 !important;
        box-shadow: 0 0 0 3px rgba(31,41,55,.1) !important;
        outline: none !important;
    }
    .stTextInput input::placeholder { color: #94a3b8 !important; }
    .stTextInput label {
        font-size: 0.78rem !important;
        font-weight: 700 !important;
        color: #374151 !important;
        letter-spacing: .04em !important;
        text-transform: uppercase !important;
    }

    /* ── Sign-in button (form_submit_button renders inside .stFormSubmitButton) ── */
    [data-testid="stMainBlockContainer"] .stButton > button[kind="primary"],
    [data-testid="stMainBlockContainer"] .stFormSubmitButton > button,
    [data-testid="stMainBlockContainer"] .stFormSubmitButton > button[kind="primary"] {
        background: linear-gradient(135deg, #1f2937 0%, #374151 100%) !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 11px !important;
        padding: 14px 0 !important;
        font-size: .95rem !important;
        font-weight: 700 !important;
        letter-spacing: .035em !important;
        box-shadow: 0 4px 18px rgba(31,41,55,.38) !important;
        transition: all .22s ease !important;
    }
    [data-testid="stMainBlockContainer"] .stButton > button[kind="primary"]:hover,
    [data-testid="stMainBlockContainer"] .stFormSubmitButton > button:hover {
        background: linear-gradient(135deg, #111827 0%, #1f2937 100%) !important;
        color: #ffffff !important;
        box-shadow: 0 10px 30px rgba(0,0,0,.35) !important;
        transform: translateY(-2px) !important;
    }
    [data-testid="stMainBlockContainer"] .stButton > button[kind="primary"]:active,
    [data-testid="stMainBlockContainer"] .stFormSubmitButton > button:active {
        transform: translateY(0) !important;
        box-shadow: 0 4px 12px rgba(0,0,0,.25) !important;
    }
    /* Ensure text inside the button is always white */
    [data-testid="stMainBlockContainer"] .stFormSubmitButton > button p,
    [data-testid="stMainBlockContainer"] .stFormSubmitButton > button span {
        color: #ffffff !important;
    }

    /* ── Alerts on login page ── */
    .stAlert {
        border-radius: 10px !important;
        font-size: .85rem !important;
        font-family: 'Inter', sans-serif !important;
    }
    .stAlert p, .stAlert span { font-size: .85rem !important; color: inherit !important; }
    </style>
    """, unsafe_allow_html=True)

    # ── Accent bar (top of card) ──────────────────────────────────────────
    st.markdown("""<div style="height:5px;margin:0 -32px;background:linear-gradient(90deg,#1f2937 0%,#6366f1 35%,#8b5cf6 55%,#6366f1 75%,#1f2937 100%);background-size:200% 100%;animation:shimmer 3s linear infinite;"></div>""", unsafe_allow_html=True)

    # ── Hero section ─────────────────────────────────────────────────────
    if os.path.exists(LOGO_PATH):
        img_b64 = get_base64_image(LOGO_PATH)
        logo_html = f'<img src="data:image/png;base64,{img_b64}" style="height:54px;width:auto;margin-bottom:16px;filter:drop-shadow(0 4px 12px rgba(0,0,0,.12));">'
    else:
        logo_html = '<div style="font-size:1.6rem;font-weight:800;color:#0f172a;letter-spacing:-.03em;margin-bottom:16px;">42Signals</div>'

    st.markdown(f"""<div style="text-align:center;padding:36px 32px 24px 32px;margin:0 -32px;background:linear-gradient(180deg,#f8fafc 0%,#ffffff 100%);border-bottom:1px solid #e2e8f0;">
{logo_html}
<div style="font-size:1.55rem;font-weight:800;color:#0f172a;letter-spacing:-.035em;line-height:1.15;font-family:'Inter',sans-serif;">Welcome back</div>
<div style="font-size:0.875rem;color:#475569;margin-top:9px;font-family:'Inter',sans-serif;line-height:1.55;font-weight:400;">Sign in to access the 42Signals<br>Requirement Handling portal</div>
<div style="display:flex;justify-content:center;gap:8px;margin-top:18px;flex-wrap:wrap;">
<span style="background:#f1f5f9;border:1px solid #e2e8f0;border-radius:20px;padding:4px 12px;font-size:0.7rem;color:#475569;font-weight:600;font-family:'Inter',sans-serif;">&#128203; Forms</span>
<span style="background:#f1f5f9;border:1px solid #e2e8f0;border-radius:20px;padding:4px 12px;font-size:0.7rem;color:#475569;font-weight:600;font-family:'Inter',sans-serif;">&#128202; Feasibility</span>
<span style="background:#f1f5f9;border:1px solid #e2e8f0;border-radius:20px;padding:4px 12px;font-size:0.7rem;color:#475569;font-weight:600;font-family:'Inter',sans-serif;">&#128256; Workflows</span>
</div>
</div>""", unsafe_allow_html=True)

    # ── Lockout check (server-side — persists across tabs) ───────────────
    # We need a username to check server-side lockout, so show the field first
    # for the lockout lookup. We re-read it after form submission too.
    _preview_user = st.session_state.get("login_username_preview", "")

    now = time.time()
    _srv_attempts, _srv_lockout_until = _get_lockout(_preview_user) if _preview_user else (0, 0.0)
    # Also check session-state lockout (covers anonymous pre-username-entry state)
    _ss_lockout = st.session_state["lockout_until"]
    locked    = (_srv_lockout_until > now) or (_ss_lockout > now)
    remaining = int(max(_srv_lockout_until, _ss_lockout) - now)

    if locked:
        st.markdown(f"""<div role="alert" style="background:linear-gradient(135deg,#fef2f2,#fee2e2);border:1px solid #fca5a5;border-left:4px solid #dc2626;border-radius:12px;padding:15px 18px;margin-bottom:4px;color:#7f1d1d;font-family:'Inter',sans-serif;font-size:.875rem;display:flex;align-items:center;gap:12px;">
<span style="font-size:1.5rem;flex-shrink:0;">&#128274;</span>
<div><div style="font-weight:700;margin-bottom:3px;">Account temporarily locked</div>
<div style="opacity:.8;">Too many failed attempts.<br>Try again in <strong>{remaining // 60}m {remaining % 60}s</strong>.</div>
</div></div>""", unsafe_allow_html=True)
        st.markdown("""<div style="text-align:center;padding:20px 0 6px 0;font-size:.72rem;color:#94a3b8;font-family:'Inter',sans-serif;">Access restricted to authorised users only &middot; 42Signals &copy; 2026</div>""", unsafe_allow_html=True)
        # Auto-refresh every second so the countdown ticks live
        time.sleep(1)
        st.rerun()

    # ── Form (st.form enables Enter-to-submit) ────────────────────────────
    with st.form("login_form", clear_on_submit=False):
        username = st.text_input(
            "Username",
            placeholder="e.g. shanjai",
            key="login_username",
            autocomplete="username",
        )
        password = st.text_input(
            "Password",
            type="password",
            placeholder="••••••••••",
            key="login_password",
            autocomplete="current-password",
        )

        # Failed-attempt inline alert (inside form so it's visible before submit)
        attempts = st.session_state.get("failed_attempts", 0)
        if attempts > 0:
            left = MAX_ATTEMPTS - attempts
            st.markdown(f"""<div role="alert" style="background:linear-gradient(135deg,#fffbeb,#fef9ec);border:1px solid #fcd34d;border-left:4px solid #f59e0b;border-radius:10px;padding:11px 15px;color:#78350f;font-size:.82rem;margin-top:4px;font-family:'Inter',sans-serif;display:flex;align-items:center;gap:9px;">
<span style="font-size:1rem;flex-shrink:0;">&#9888;&#65039;</span>
<span>Incorrect credentials &mdash; <strong>{left} attempt{'s' if left != 1 else ''}</strong> left before lockout.</span>
</div>""", unsafe_allow_html=True)

        st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)
        submitted = st.form_submit_button("Sign In  →", type="primary", use_container_width=True)

    if submitted:
        clean_user = username.strip().lower()
        # Store username preview for server-side lockout lookup on next render
        st.session_state["login_username_preview"] = clean_user
        # Re-check server-side lockout for this specific username
        _srv_attempts, _srv_lockout_until = _get_lockout(clean_user)
        if _srv_lockout_until > time.time():
            st.rerun()
        elif not username or not password:
            st.warning("Please enter both username and password.")
        elif verify_password(username, password):
            user = get_user(username)
            display_name = user["display_name"] if user else clean_user
            st.session_state["authenticated"]   = True
            st.session_state["current_user"]    = clean_user
            st.session_state["display_name"]    = display_name
            st.session_state["failed_attempts"] = 0
            st.session_state["lockout_until"]   = 0.0
            _clear_lockout(clean_user)
            _new_analytics_sid = str(uuid.uuid4())
            st.session_state["analytics_sid"] = _new_analytics_sid
            log_event(EVENT_LOGIN, clean_user, _new_analytics_sid)
            _sid = _save_session(clean_user, display_name)
            st.session_state["ls_write_token"] = _sid
            st.rerun()
        else:
            new_attempts = _srv_attempts + 1
            if new_attempts >= MAX_ATTEMPTS:
                _set_lockout(clean_user, 0, time.time() + LOCKOUT_SECONDS)
                st.session_state["lockout_until"]   = time.time() + LOCKOUT_SECONDS
                st.session_state["failed_attempts"] = 0
            else:
                _set_lockout(clean_user, new_attempts, 0.0)
                st.session_state["failed_attempts"] = new_attempts
            st.rerun()

    # ── Card footer ──────────────────────────────────────────────────────
    st.markdown("""<div style="text-align:center;padding:20px 32px 6px 32px;margin:8px -32px 0 -32px;border-top:1px solid #f1f5f9;">
<div style="font-size:.72rem;color:#94a3b8;font-family:'Inter',sans-serif;display:flex;align-items:center;justify-content:center;gap:10px;">
<span style="display:inline-block;width:24px;height:1px;background:#e2e8f0;"></span>
Access restricted to authorised users only
<span style="display:inline-block;width:24px;height:1px;background:#e2e8f0;"></span>
</div></div>""", unsafe_allow_html=True)

    # ── Below-card caption ────────────────────────────────────────────────
    st.markdown("""<div style="text-align:center;margin-top:22px;font-size:.7rem;color:rgba(255,255,255,.2);font-family:'Inter',sans-serif;letter-spacing:.05em;">42Signals &nbsp;&copy;&nbsp; 2026</div>""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR  (CSS-hidden on login page via render_login(); Python still runs safely)
# ─────────────────────────────────────────────────────────────────────────────

with st.sidebar:

    # Brand header
    if os.path.exists(LOGO_PATH):
        img_b64 = get_base64_image(LOGO_PATH)
        st.markdown(f"""
        <div style="text-align:center; padding:28px 16px 20px 16px;">
            <img src="data:image/png;base64,{img_b64}"
                 style="height:52px; width:auto; margin-bottom:12px; display:block; margin-left:auto; margin-right:auto;">
            <div style="font-size:0.7rem; font-weight:600; color:#9ca3af; letter-spacing:0.12em; text-transform:uppercase; font-family:'Inter',sans-serif;">Requirement Handling</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="text-align:center; padding:28px 16px 20px 16px;">
            <div style="font-size:1.1rem; font-weight:700; color:#1f2937;">42Signals</div>
            <div style="font-size:0.7rem; font-weight:600; color:#9ca3af; letter-spacing:0.12em; text-transform:uppercase; margin-top:4px;">Requirement Handling</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<hr style="border:none;border-top:1px solid #e5e7eb;margin:0 0 14px 0;">', unsafe_allow_html=True)

    # Section label
    st.markdown('<div style="color:#6b7280;font-size:0.67rem;text-transform:uppercase;letter-spacing:0.14em;padding:0 6px 10px 6px;font-weight:600;font-family:\'Inter\',sans-serif;">Navigation</div>', unsafe_allow_html=True)

    # Nav items — SVG icons, consistent active/inactive styling
    _NAV = {
        "main": (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><rect x="4" y="2" width="16" height="20" rx="2"/><path d="M8 7h8M8 11h8M8 15h5"/></svg>',
            "New Requirement Form",
        ),
        "feasibility": (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><path d="M3 20h18M7 20V12M12 20V5M17 20v-8"/></svg>',
            "Feasibility Assessment",
        ),
        "req_flow": (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><circle cx="5" cy="12" r="2.5"/><circle cx="19" cy="12" r="2.5"/><path d="M7.5 12h9"/><path d="M14.5 9l3 3-3 3"/></svg>',
            "New Requirement Flow",
        ),
        "ops_map": (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="7" height="7" rx="1"/><rect x="14" y="3" width="7" height="7" rx="1"/><rect x="3" y="14" width="7" height="7" rx="1"/><rect x="14" y="14" width="7" height="7" rx="1"/></svg>',
            "Day-to-Day Ops Map",
        ),
        "poc_guide": (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="8" r="4"/><path d="M4 20c0-4 3.6-7 8-7s8 3 8 7"/></svg>',
            "Task POC Guide",
        ),
        "cost_calc": (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="18" height="18" rx="2"/><path d="M7 8h2m4 0h3M7 12h2m4 0h3M7 16h2m4 0h3"/></svg>',
            "Cost Calculator",
        ),
        "ext_tools": (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><path d="M10 13a5 5 0 0 0 7.54.54l3-3a5 5 0 0 0-7.07-7.07l-1.72 1.71"/><path d="M14 11a5 5 0 0 0-7.54-.54l-3 3a5 5 0 0 0 7.07 7.07l1.71-1.71"/></svg>',
            "External Tools",
        ),
    }

    # Analytics Dashboard — admin-only nav entry
    _current_role = (get_user(st.session_state.get("current_user", "") or "") or {}).get("role", "")
    if _current_role == "admin":
        _NAV["analytics"] = (
            '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><path d="M3 20h18M7 20V12M12 20V5M17 20v-8"/></svg>',
            "Analytics Dashboard",
        )

    for key, (svg, label) in _NAV.items():
        active = st.session_state["page"] == key
        if active:
            st.markdown(
                f"""<div style="background:linear-gradient(135deg,#f1f5f9 0%,#e8ecf0 100%);border:1px solid #dde3ea;border-left:3px solid #1f2937;border-radius:8px;padding:9px 12px;color:#111827;font-size:0.83rem;font-weight:600;margin-bottom:3px;display:flex;align-items:center;gap:9px;font-family:'Inter',sans-serif;white-space:nowrap;box-shadow:0 1px 3px rgba(0,0,0,0.05);">{svg}&nbsp;{label}</div>""",
                unsafe_allow_html=True,
            )
        else:
            if st.button(label, key=f"nav_{key}", use_container_width=True):
                st.session_state["page"] = key
                st.rerun()

    st.markdown('<hr style="border:none;border-top:1px solid #e5e7eb;margin:14px 0 10px 0;">', unsafe_allow_html=True)

    # ── Logged-in user info + logout ──────────────────────────────────────────
    display_name = st.session_state.get("display_name") or ""
    st.markdown(f"""
    <div style="
        background: linear-gradient(135deg,#f8fafc 0%,#f1f5f9 100%);
        border:1px solid #e5e7eb; border-radius:10px;
        padding:11px 14px; margin-bottom:10px;
        font-family:'Inter',sans-serif;
        display:flex; align-items:center; gap:10px;
    ">
        <div style="
            width:32px; height:32px; border-radius:50%;
            background:linear-gradient(135deg,#1f2937 0%,#374151 100%);
            display:flex; align-items:center; justify-content:center;
            font-size:0.875rem; font-weight:700; color:#fff; flex-shrink:0;
        ">{_h(display_name[:1].upper()) if display_name else "?"}</div>
        <div>
            <div style="font-size:0.82rem;font-weight:600;color:#111827;">{_h(display_name)}</div>
            <div style="font-size:0.7rem;color:#6b7280;margin-top:1px;">Signed in</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Sign Out", key="logout_btn", use_container_width=True):
        log_event(EVENT_LOGOUT, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""))
        _clear_session()
        st.query_params.clear()
        st.session_state["ls_clear"]        = True
        st.session_state["authenticated"]   = False
        st.session_state["current_user"]    = None
        st.session_state["display_name"]    = None
        st.session_state["failed_attempts"] = 0
        st.session_state["lockout_until"]   = 0.0
        st.session_state["page"]                = "main"
        st.session_state["cc_show_results"]     = False
        st.session_state["analytics_sid"]       = str(uuid.uuid4())
        st.session_state["analytics_last_page"] = ""
        st.rerun()

    st.markdown("""
    <div style="text-align:center; padding:6px 0 12px 0;">
        <div style="color:#9ca3af; font-size:0.68rem; font-family:'Inter',sans-serif; letter-spacing:0.04em;">v1.0 &nbsp;·&nbsp; 42Signals &nbsp;·&nbsp; 2026</div>
    </div>
    """, unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# SUBMISSION PERSISTENCE
# ─────────────────────────────────────────────────────────────────────────────

_SUBMISSIONS_DIR = Path("submissions")
_FORM_KEY_PREFIXES = (
    "form_", "pt_", "sos_", "rev_", "pv_", "storeid_", "festive_", "final_",
)


def _json_default(obj):
    """JSON serialiser for types not handled natively (date, set, etc.)."""
    if isinstance(obj, date):
        return obj.isoformat()
    if isinstance(obj, set):
        return list(obj)
    return str(obj)


def save_submission(form_data: dict, client_name: str, username: str) -> None:
    """Persist the current form widget state and form_data to submissions/.
    If a submission was loaded for editing, overwrite that file instead of creating a new one."""
    _SUBMISSIONS_DIR.mkdir(exist_ok=True)
    editing_file = st.session_state.get("_editing_submission_file")
    if editing_file and (_SUBMISSIONS_DIR / editing_file).exists():
        filename = editing_file
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", client_name or "unknown")
        filename = f"{safe_name}_{timestamp}.json"
    # Capture all widget keys that belong to the requirement form
    snapshot = {
        k: v for k, v in st.session_state.items()
        if isinstance(k, str) and k.startswith(_FORM_KEY_PREFIXES)
    }
    payload = {
        "client_name": client_name,
        "saved_at": datetime.now().isoformat(),
        "saved_by": username,
        "form_data": form_data,
        "session_state": snapshot,
    }
    with open(_SUBMISSIONS_DIR / filename, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, indent=2, default=_json_default)
    st.session_state["_editing_submission_file"] = filename
    list_submissions.clear()  # invalidate cache so the new file appears immediately


@st.cache_data(ttl=60)
def list_submissions() -> list[dict]:
    """Return list of submission metadata dicts, newest first.
    Cached for 60 s so repeated reruns don't re-read the filesystem."""
    if not _SUBMISSIONS_DIR.exists():
        return []
    result = []
    for p in sorted(_SUBMISSIONS_DIR.glob("*.json"), reverse=True):
        try:
            with open(p, encoding="utf-8") as fh:
                data = json.load(fh)
            result.append({
                "filename": p.name,
                "client_name": data.get("client_name", p.stem),
                "saved_at": data.get("saved_at", ""),
                "saved_by": data.get("saved_by", ""),
            })
        except Exception:
            pass
    return result


def load_submission(filename: str) -> None:
    """Restore form widget state from a saved submission JSON and trigger rerun."""
    path = _SUBMISSIONS_DIR / filename
    if not path.exists():
        st.error("Submission file not found.")
        return
    with open(path, encoding="utf-8") as fh:
        data = json.load(fh)
    _ISO_DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
    for k, v in data.get("session_state", {}).items():
        # Re-hydrate any ISO date string back to datetime.date.
        # Per-domain date keys get a suffix (e.g. pt_inputs_expected_date_swiggy_com)
        # so key-name matching is unreliable — match on value shape instead.
        if isinstance(v, str) and _ISO_DATE_RE.match(v):
            try:
                v = date.fromisoformat(v)
            except ValueError:
                pass
        st.session_state[k] = v
    st.session_state["_editing_submission_file"] = filename
    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: New Requirement Form
# ─────────────────────────────────────────────────────────────────────────────

def _field_filled(section, field):
    """True if field is non-empty directly, or via any per-domain key ('domain — field')."""
    val = section.get(field)
    if isinstance(val, list):
        return bool(val)
    if val:
        return True
    return any(v for k, v in section.items() if k.endswith(f"\u2014 {field}") and v)


def _validate_form(form_data, modules):
    """Return list of error strings for all missing required fields."""
    errors = []

    if not form_data.get("Client Information", {}).get("Target Market"):
        errors.append("**Client Information** \u2014 Target Market / Geography is required.")

    if not modules:
        errors.append("**Modules** \u2014 Select at least one module.")

    if "Products + Trends" in modules:
        if not form_data.get("Products + Trends", {}).get("Domains"):
            errors.append("**Products + Trends** \u2014 At least one domain is required.")

    if "SOS (Search on Site)" in modules:
        sos = form_data.get("SOS (Search on Site)", {})
        if not sos.get("Domains"):
            errors.append("**SOS** \u2014 At least one domain is required.")
        if not int(sos.get("No. of Keywords", 0)):
            errors.append("**SOS** \u2014 Number of Keywords must be greater than 0.")

    if "Reviews" in modules:
        rev = form_data.get("Reviews", {})
        if not rev.get("Domains"):
            errors.append("**Reviews** \u2014 At least one domain is required.")
        if not _field_filled(rev, "Input Sources"):
            errors.append("**Reviews** \u2014 At least one input source must be selected.")

    if "Price Violation" in modules:
        pv = form_data.get("Price Violation", {})
        if not pv.get("Domains"):
            errors.append("**Price Violation** \u2014 At least one domain is required.")
        if not _field_filled(pv, "Product URL List"):
            errors.append("**Price Violation** \u2014 Product URL list is required.")

    if "Store ID Crawls" in modules:
        if not form_data.get("Store ID Crawls", {}).get("Domains"):
            errors.append("**Store ID Crawls** \u2014 At least one domain is required.")

    if "Festive Sale Crawls" in modules:
        festive = form_data.get("Festive Sale Crawls", {})
        if festive.get("Crawl Type") == "Products + Trends Based" and not festive.get("Domains"):
            errors.append("**Festive Sale Crawls** \u2014 Domains required for Products + Trends Based type.")

    if not form_data.get("Final Alignment", {}).get("Client Core Objective"):
        errors.append("**Final Alignment** \u2014 Client Core Objective is required.")

    return errors


def render_main_form():
    page_title(
        "New Requirement Form",
        "Capture complete client crawl requirements for project planning and scoping."
    )

    # ── Load Previous Submission ───────────────────────────────────────────
    _saved = list_submissions()
    if _saved:
        with st.expander("📂  Load a previous submission", expanded=False):
            _options = {
                f"{s['client_name']}  ·  {s['saved_at'][:16].replace('T', ' ')}  (by {s['saved_by']})": s["filename"]
                for s in _saved
            }
            _chosen_label = st.selectbox(
                "Select submission to load",
                list(_options.keys()),
                key="_load_submission_select",
                label_visibility="collapsed",
            )
            btn_col1, _ = st.columns([1, 3])
            with btn_col1:
                if st.button("⬆️  Load & Edit", key="_load_submission_btn"):
                    load_submission(_options[_chosen_label])

    # If editing a loaded submission, show a banner + allow starting fresh
    if st.session_state.get("_editing_submission_file"):
        info_col, btn_col = st.columns([4, 1])
        with info_col:
            st.info(f"✏️  Editing: **{st.session_state['_editing_submission_file']}** — saving will update this record in place.")
        with btn_col:
            if st.button("✚  New Form", key="_new_form_btn", use_container_width=True):
                st.session_state["_editing_submission_file"] = None
                for k in list(st.session_state.keys()):
                    if isinstance(k, str) and k.startswith(_FORM_KEY_PREFIXES):
                        del st.session_state[k]
                st.rerun()

    left, right = st.columns([2, 1], gap="large")
    form_data = {}

    with left:

        st.markdown('<div style="font-size:0.75rem;color:#6b7280;margin-bottom:8px;"><span style="color:#dc2626;font-weight:700;">*</span> Required field</div>', unsafe_allow_html=True)

        # ── 1. Client Information ──────────────────────────────────────────
        section_header("👤", "1. Client Information")

        c1, c2 = st.columns([2, 1])
        with c1:
            client_name = st.text_input("Client Name *", placeholder="e.g., Unilever India", key="form_client_name")
        with c2:
            priority = st.selectbox("Priority Level", ["High", "Medium", "Low"], key="form_priority")

        c3, c4 = st.columns(2)
        with c3:
            completion_date = st.date_input("Expected Completion Date", key="form_completion_date")
        with c4:
            expected_market = st.text_input("Target Market / Geography *", placeholder="e.g., India, Southeast Asia", key="form_target_market")

        form_data["Client Information"] = {
            "Client Name":            client_name,
            "Priority Level":         priority,
            "Expected Completion":    str(completion_date),
            "Target Market":          expected_market,
        }

        validate_required(client_name)

        # ── 2. Modules to Crawl ───────────────────────────────────────────
        section_header("🧩", "2. Modules to Crawl")

        modules = st.multiselect(
            "Select the modules required for this client *",
            ["Products + Trends", "SOS (Search on Site)", "Reviews",
             "Price Violation", "Store ID Crawls", "Festive Sale Crawls"],
            key="form_modules",
        )
        if not modules:
            st.info("Select at least one module above to configure its settings.")
        form_data["Modules Selected"] = {
            "Selected Modules": ", ".join(modules) if modules else "None"
        }

        # ── 3. Products + Trends ──────────────────────────────────────────
        if "Products + Trends" in modules:
            section_header("📦", "3. Products + Trends Module")
            pt = {}

            pt["Domains"], _pt_domains = domain_selector("Domains *", "pt")

            _apply_domain_config(pt, "pt", _pt_domains, _pt_crawl_config)

            st.markdown("**Specific Fields to Capture**")
            pt["Specific Fields"] = st.text_area(
                "Any additional fields to extract",
                placeholder="e.g., seller name, discount %, stock status, rating breakdown, delivery time",
                key="pt_specific_fields",
            )
            form_data["Products + Trends"] = pt

        # ── 4. SOS Module ─────────────────────────────────────────────────
        if "SOS (Search on Site)" in modules:
            section_header("🔍", "4. SOS (Search On Site) Module")
            sos = {}

            st.markdown("**Keywords**")
            c1, c2 = st.columns(2)
            with c1:
                sos["No. of Keywords"] = st.number_input("Number of Keywords *", min_value=0, key="sos_keyword_count")
            with c2:
                keywords_source = st.radio("Keywords source", ["Client Provided", "Provide Sample for Testing"], key="sos_keywords_source")
            if keywords_source == "Client Provided":
                sos["SOS Keywords Sheet Link"] = st.text_input("Link to client keywords sheet", key="sos_keywords_sheet_link")
            else:
                sos["Sample Keywords"] = st.text_area("Sample keywords for testing", placeholder="e.g., laptop, shoes, home appliances", key="sos_sample_keywords")

            sos["Domains"], _sos_domains = domain_selector("Domains *", "sos")

            _apply_domain_config(sos, "sos", _sos_domains, _sos_crawl_config)
            form_data["SOS (Search on Site)"] = sos

        # ── 5. Reviews Module ─────────────────────────────────────────────
        if "Reviews" in modules:
            section_header("⭐", "5. Reviews Module")
            rev = {}

            rev["Domains"], _rev_domains = domain_selector("Domains *", "reviews")

            _apply_domain_config(rev, "rev", _rev_domains, _rev_crawl_config)
            form_data["Reviews"] = rev

        # ── 6. Price Violation Module ─────────────────────────────────────
        if "Price Violation" in modules:
            section_header("💰", "6. Price Violation Module")
            pv = {}

            pv["Domains"], _pv_domains = domain_selector("Domains *", "pv")

            _apply_domain_config(pv, "pv", _pv_domains, _pv_crawl_config)
            form_data["Price Violation"] = pv

        # ── 7. Store ID Crawls ────────────────────────────────────────────
        if "Store ID Crawls" in modules:
            section_header("🏪", "7. Store ID Crawl")
            storeid = {}

            storeid["Domains"], _storeid_domains = domain_selector("Domains *", "storeid")

            _apply_domain_config(storeid, "storeid", _storeid_domains, _storeid_crawl_config)
            form_data["Store ID Crawls"] = storeid

        # ── 8. Festive Sale Crawls ────────────────────────────────────────
        if "Festive Sale Crawls" in modules:
            section_header("🎉", "8. Festive Sale Crawls")
            festive = {}

            festive["Crawl Type"] = st.radio(
                "Crawl Type",
                ["Products + Trends Based", "SOS Type", "Category URL Based"],
                key="festive_type", horizontal=True,
            )
            if festive["Crawl Type"] == "Products + Trends Based":
                festive["Domains"], _festive_domains = domain_selector("Domains *", "festive")
                st.markdown("**Schedule**")
                _apply_domain_config(festive, "festive", _festive_domains, _festive_schedule_config)
            elif festive["Crawl Type"] == "Category URL Based":
                festive["Category URL List"] = st.text_area("Category URLs", placeholder="Provide category URLs for festive crawl", key="festive_category_urls")
                st.markdown("**Schedule**")
                festive.update(_festive_schedule_config())
            else:
                st.markdown("**Schedule**")
                festive.update(_festive_schedule_config())
            form_data["Festive Sale Crawls"] = festive

        # ── 9. Final Alignment ────────────────────────────────────────────
        section_header("🎯", "9. Final Alignment")
        form_data["Final Alignment"] = {
            "Client Core Objective": st.text_area(
                "What is the client's core objective? *",
                placeholder="e.g., Market gap analysis, brand monitoring, competitive pricing intelligence, demand trends…",
                key="final_objective",
            ),
            "Expectations From Us": st.text_area(
                "What are you expecting from us?",
                placeholder="e.g., Real-time dashboards, daily reports, anomaly alerts, competitive benchmarking…",
                key="final_expectation",
            ),
        }

        # ── 10. Comments ──────────────────────────────────────────────────
        section_header("💬", "10. Comments & Notes")
        form_data["Comments & Notes"] = {
            "Additional Comments": st.text_area(
                "Any additional notes or special instructions",
                placeholder="Anything else important for the team to know…",
                key="final_comments",
            )
        }

        st.markdown("<br>", unsafe_allow_html=True)

        # PDF Generation
        if st.button("⬇️  Generate & Download PDF", type="primary", use_container_width=True):
            if not client_name:
                st.error("Enter a Client Name before generating the PDF.")
                st.stop()
            _errors = _validate_form(form_data, modules)
            if _errors:
                st.error("Please fix the following before generating the PDF:")
                for _e in _errors:
                    st.markdown(f"- {_e}")
                st.stop()
            log_event(EVENT_GENERATE_REQ_PDF, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "main", {"client": client_name})
            try:
                with st.spinner("Building PDF…"):
                    pdf_bytes = generate_pdf(form_data, client_name).read()
                st.session_state["pdf_bytes"] = pdf_bytes
            except Exception as e:
                st.error(f"PDF generation failed — please try again or contact your admin. ({type(e).__name__})")
                st.stop()
            # Save form state so it can be reloaded and edited later
            try:
                save_submission(form_data, client_name, st.session_state.get("current_user", ""))
            except Exception:
                pass  # never block PDF generation due to a save error
            # _safe_filename strips path-traversal chars from user-supplied client name
            st.session_state["pdf_name"] = _safe_filename(client_name, "_Requirement_Form.pdf")
            celebrate(
                message="PDF generated successfully!",
                sub=f"{_h(client_name)} Requirement Form is ready to download."
            )
            st.toast("PDF ready! Click below to download.", icon="🎉")
            components.html(
                "<script>window.parent.document.querySelector('[data-testid=\"stAppViewContainer\"] > section')?.scrollTo({top:999999,behavior:'smooth'});</script>",
                height=0,
            )

        if st.session_state.get("pdf_bytes"):
            if st.download_button(
                label="📄  Download Requirement PDF",
                data=st.session_state["pdf_bytes"],
                file_name=st.session_state.get("pdf_name", "requirement.pdf"),
                mime="application/pdf",
                use_container_width=True,
            ):
                log_event(EVENT_DOWNLOAD_REQ_PDF, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "main")

    with right:
        render_summary(form_data)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: FEASIBILITY ASSESSMENT
# ─────────────────────────────────────────────────────────────────────────────

def render_feasibility():
    # Clear any previously generated doc when the page is freshly visited
    # so stale downloads from earlier sessions don't persist.
    if st.session_state.get("analytics_last_page") != "feasibility":
        st.session_state.pop("feas_doc", None)
        st.session_state.pop("feas_name", None)

    page_title(
        "Feasibility Assessment",
        "Evaluate crawl feasibility before project kickoff — exports a Word document for the team."
    )

    # Stats strip
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("""
        <div style="background:white;border-radius:12px;padding:18px 20px;border-left:4px solid #1f2937;box-shadow:0 2px 8px rgba(0,0,0,0.07);transition:box-shadow 0.2s;">
            <div style="font-size:0.67rem;color:#94a3b8;text-transform:uppercase;letter-spacing:0.09em;font-weight:700;font-family:'Inter',sans-serif;">Purpose</div>
            <div style="font-size:0.9rem;font-weight:600;color:#111827;margin-top:5px;font-family:'Inter',sans-serif;">Pre-project scoping</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown("""
        <div style="background:white;border-radius:12px;padding:18px 20px;border-left:4px solid #1f2937;box-shadow:0 2px 8px rgba(0,0,0,0.07);transition:box-shadow 0.2s;">
            <div style="font-size:0.67rem;color:#94a3b8;text-transform:uppercase;letter-spacing:0.09em;font-weight:700;font-family:'Inter',sans-serif;">Output</div>
            <div style="font-size:0.9rem;font-weight:600;color:#111827;margin-top:5px;font-family:'Inter',sans-serif;">Word Document (.docx)</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown("""
        <div style="background:white;border-radius:12px;padding:18px 20px;border-left:4px solid #1f2937;box-shadow:0 2px 8px rgba(0,0,0,0.07);transition:box-shadow 0.2s;">
            <div style="font-size:0.67rem;color:#94a3b8;text-transform:uppercase;letter-spacing:0.09em;font-weight:700;font-family:'Inter',sans-serif;">Use Case</div>
            <div style="font-size:0.9rem;font-weight:600;color:#111827;margin-top:5px;font-family:'Inter',sans-serif;">Share with tech / ops team</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    col, _ = st.columns([3, 1])
    with col:

        # Requirement Information
        section_header("📌", "Requirement Information")
        c1, c2 = st.columns(2)
        with c1:
            client_name = st.text_input("Client Name", placeholder="e.g., Unilever India", key="feas_client")
        with c2:
            requestor_name = st.text_input("Requestor Name", placeholder="Your name", key="feas_requestor")

        # Domains
        section_header("🌐", "Domain List")
        num_domains = st.number_input("Number of Domains", min_value=1, max_value=50, step=1, value=1, key="feas_num_domains")
        domains = []
        if num_domains <= 6:
            cols = st.columns(min(int(num_domains), 3))
            for i in range(int(num_domains)):
                with cols[i % 3]:
                    d = st.text_input(f"Domain {i+1}", placeholder="example.com", key=f"feas_domain_{i}")
                    if d:
                        domains.append(d)
        else:
            for i in range(int(num_domains)):
                d = st.text_input(f"Domain {i+1}", placeholder="example.com", key=f"feas_domain_{i}")
                if d:
                    domains.append(d)

        # Crawl Configuration
        section_header("⚙️", "Crawl Configuration")
        crawl_options = st.multiselect(
            "Select crawl type and special requirements",
            ["Category Based", "Product URL Input Based", "SOS", "Reviews",
             "Festive Sales Day Crawl", "Banner Crawl", "Others"],
            key="feas_crawl_options",
        )
        crawl_type = None
        if "Category Based" in crawl_options:
            crawl_type = "Category Based"
        elif "Product URL Input Based" in crawl_options:
            crawl_type = "Product URL Input Based"

        crawl_features = [o for o in crawl_options if o not in ["Category Based", "Product URL Input Based"]]
        others_desc = ""
        if "Others" in crawl_features:
            others_desc = st.text_input("If Others, please specify", key="feas_others")

        # Zipcode
        section_header("📍", "Zipcode Requirement")
        zipcode_type = st.radio(
            "Zipcode handling", ["With Zipcode", "Without Zipcode", "Both"],
            horizontal=True, key="feas_zipcode"
        )
        target_city = target_state = target_country = ""
        if zipcode_type in ["With Zipcode", "Both"]:
            st.markdown("**Target Location**")
            c1, c2, c3 = st.columns(3)
            with c1:
                target_city = st.text_input("City", key="feas_city")
            with c2:
                target_state = st.text_input("State", key="feas_state")
            with c3:
                target_country = st.text_input("Country", key="feas_country")

        # Additional
        section_header("📝", "Additional Information")
        additional_notes = st.text_area("Additional details / notes", key="feas_notes", height=120)

        st.markdown("<br>", unsafe_allow_html=True)

        # Generate document
        if st.button("📄  Generate Feasibility Document", type="primary", use_container_width=True):
            if not client_name:
                st.error("Enter a Client Name before generating the document.")
                st.stop()
            else:
                log_event(EVENT_GENERATE_FEAS, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "feasibility", {"client": client_name})
                try:
                    with st.spinner("Building document…"):
                        from docx import Document  # type: ignore
                        doc = Document()
                        doc.add_heading(f"{client_name} — Feasibility Requirement", level=1)

                        doc.add_heading("Requirement Information", level=2)
                        doc.add_paragraph(f"Client Name: {client_name}")
                        doc.add_paragraph(f"Requestor Name: {requestor_name}")

                        doc.add_heading("Domains", level=2)
                        for d in domains:
                            doc.add_paragraph(d, style="List Bullet")

                        doc.add_heading("Crawl Configuration", level=2)
                        doc.add_paragraph(f"Crawl Type: {crawl_type or 'Not specified'}")
                        doc.add_paragraph(f"Special Requirements: {', '.join(crawl_features) or 'None'}")
                        if others_desc:
                            doc.add_paragraph(f"Others: {others_desc}")

                        doc.add_heading("Zipcode Requirement", level=2)
                        doc.add_paragraph(f"Zipcode Handling: {zipcode_type}")
                        if zipcode_type in ["With Zipcode", "Both"]:
                            doc.add_heading("Target Location", level=2)
                            doc.add_paragraph(f"City: {target_city}")
                            doc.add_paragraph(f"State: {target_state}")
                            doc.add_paragraph(f"Country: {target_country}")

                        doc.add_heading("Additional Notes", level=2)
                        doc.add_paragraph(additional_notes or "None")

                        buf = BytesIO()
                        doc.save(buf)
                        buf.seek(0)

                    st.session_state["feas_doc"]  = buf.getvalue()
                    st.session_state["feas_name"] = _safe_filename(client_name, "_Feasibility_Requirement.docx")
                    celebrate(
                        message="Feasibility Document generated!",
                        sub=f"{_h(client_name)} feasibility doc is ready to download."
                    )
                    st.toast("Document ready! Click below to download.", icon="🎉")
                except Exception as e:
                    st.error(f"Document generation failed — please try again or contact your admin. ({type(e).__name__})")

        if st.session_state.get("feas_doc"):
            if st.download_button(
                label="⬇️  Download Feasibility Document",
                data=st.session_state["feas_doc"],
                file_name=st.session_state.get("feas_name", "feasibility.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            ):
                log_event(EVENT_DOWNLOAD_FEAS, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "feasibility")


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: DECISION MIND MAP
# ─────────────────────────────────────────────────────────────────────────────

def render_req_flow():
    page_title("New Requirement Decision Tree", "Step-by-step guide from client intake to delivery. Click nodes to expand.")
    st.markdown("""<div style="background:white;border:1px solid #e5e7eb;border-radius:10px;padding:10px 20px;margin-bottom:14px;display:flex;gap:24px;flex-wrap:wrap;align-items:center;font-size:0.82rem;box-shadow:0 1px 4px rgba(0,0,0,0.04);font-family:'Inter',sans-serif;"><span style="color:#3b82f6;font-weight:600;">&#9646; Step</span><span style="color:#d97706;font-weight:600;">&#9646; Decision</span><span style="color:#22c55e;font-weight:600;">&#9646; Outcome</span><span style="color:#94a3b8;font-weight:600;">&#9646; Action</span><span style="color:#94a3b8;margin-left:auto;font-size:0.78rem;">8 sequential steps &mdash; intake to delivery</span></div>""", unsafe_allow_html=True)
    _html = """<!DOCTYPE html><meta charset="utf-8">
<style>
html,body{margin:0;padding:0;width:100vw;height:100vh;overflow:hidden;
  background:#f0f2f6;font-family:'Inter','Segoe UI',-apple-system,sans-serif;}
#tree{width:100vw;height:100vh;}
.node rect{stroke-width:1.8px;transition:all .18s;cursor:pointer;}
.node rect:hover{opacity:.82;}
.node text{pointer-events:none;font-weight:500;}
.link{fill:none;stroke-width:1.6px;opacity:.4;}
#legend{position:absolute;top:12px;right:14px;background:rgba(255,255,255,.93);
  border:1px solid #e5e7eb;border-radius:9px;padding:10px 14px;
  font-size:12px;color:#374151;line-height:2.1;
  box-shadow:0 2px 8px rgba(0,0,0,.07);}
#hint{position:absolute;bottom:10px;right:14px;font-size:10px;color:#9ca3af;}
</style>
<div id="tree"></div>
<div id="legend"><b>Node Types</b><br><span style='color:#3b82f6'>&#9646;</span> Step &nbsp;<span style='color:#d97706'>&#9646;</span> Decision &nbsp;<span style='color:#22c55e'>&#9646;</span> Outcome &nbsp;<span style='color:#94a3b8'>&#9646;</span> Action</div>
<div id="hint">Scroll = zoom &nbsp;·&nbsp; Drag = pan &nbsp;·&nbsp; Click = expand/collapse</div>
{_D3_INLINE}
<script>
var data={name:"New Requirement Intake",k:"start",children:[
{name:"1. Finalise Domains",k:"step",children:[
  {name:"QCommerce",k:"act",children:[{name:"Swiggy Instamart"},{name:"ZeptoNow"},{name:"Blinkit"},{name:"BigBasket"}]},
  {name:"ECom",k:"act",children:[{name:"Amazon"},{name:"Flipkart"},{name:"Nykaa"},{name:"Purplle"}]},
  {name:"Fashion Retail",k:"act",children:[{name:"AJIO"},{name:"Myntra"}]}
]},
{name:"2. Classify Seed URLs",k:"step",children:[
  {name:"Category URL?",k:"dec",children:[
    {name:"PDP crawl",k:"out",children:[{name:"Products Index"},{name:"Trends Index"}]},
    {name:"Listing page",k:"out",children:[{name:"SOS Index"}]},
    {name:"Banner present",k:"out",children:[{name:"Misc Index"}]}
  ]},
  {name:"Direct Product URL?",k:"dec",children:[
    {name:"Trends Index only",k:"out"}
  ]},
  {name:"SOS Keyword input?",k:"dec",children:[
    {name:"SOS Index",k:"out",children:[{name:"No PDP fetch"},{name:"Listing page only"}]}
  ]},
  {name:"Reviews needed?",k:"dec",children:[
    {name:"Reviews Index",k:"out",children:[{name:"ES input from Products"}]}
  ]},
  {name:"Banner tracking?",k:"dec",children:[
    {name:"API available",k:"out",children:[{name:"Misc API crawl"}]},
    {name:"No API",k:"out",children:[{name:"Screenshot method"}]}
  ]}
]},
{name:"3. Index Decision",k:"step",children:[
  {name:"Zipcode based?",k:"dec",children:[
    {name:"YES: add _zipcode",k:"out"},
    {name:"NO: standard name",k:"out"}
  ]},
  {name:"Historical tracking?",k:"dec",children:[
    {name:"YES: Trends/SOS/Misc",k:"out"},
    {name:"NO: Products/Reviews",k:"out"}
  ]},
  {name:"Variants needed?",k:"dec",children:[
    {name:"YES: uniq_id/variant",k:"out"},
    {name:"NO: uniq_id/product",k:"out"}
  ]},
  {name:"Assign Domain #",k:"act",children:[
    {name:"Products: no number"},
    {name:"Sheet Hourly: #1"},
    {name:"ES Daily: #2"},
    {name:"ES Hourly: #3"},
    {name:"SOS / PV: #4"},
    {name:"Others: #5+"}
  ]}
]},
{name:"4. Feasibility Check",k:"step",children:[
  {name:"Check schema sheet",k:"act",children:[
    {name:"Must-have fields"},{name:"General feasibility"}
  ]},
  {name:"Special fields?",k:"dec",children:[
    {name:"YES: define + Platform",k:"out"},
    {name:"NO: proceed",k:"out"}
  ]},
  {name:"Blocking challenge?",k:"dec",children:[
    {name:"YES: raise flag early",k:"out"},
    {name:"NO: proceed",k:"out"}
  ]}
]},
{name:"5. Site Setup",k:"step",children:[
  {name:"Domain exists?",k:"dec",children:[
    {name:"YES: reuse (even active)",k:"out"},
    {name:"NO: ask TPM to create",k:"out"}
  ]},
  {name:"Define site name",k:"act",children:[
    {name:"domain_tld{#}{_zip}"},
    {name:"_{index}_forty_two_signals"}
  ]},
  {name:"Update mapping JSON",k:"act",children:[
    {name:"domain_numbered_sites"},
    {name:"_mapping.json"}
  ]}
]},
{name:"6. Dev Implementation",k:"step",children:[
  {name:"Products: RSS→DSK→EXT",k:"act",children:[
    {name:"uniq_id unique always"},
    {name:"internal_client_name"},
    {name:"joining_key required"},
    {name:"ISO date format only"}
  ]},
  {name:"Trends: inherit DSK+EXT",k:"act",children:[
    {name:"RSS changes only"},
    {name:"uniq_id unchanged"}
  ]},
  {name:"SOS: listing only",k:"act",children:[
    {name:"No PDP fetch"},
    {name:"Scroll / page limit"}
  ]}
]},
{name:"7. Post-Setup Checks",k:"step",children:[
  {name:"Dev QA all fields",k:"act"},
  {name:"Count match?",k:"dec",children:[
    {name:"Crawlboard = Kibana",k:"out"},
    {name:"Mismatch: check logs",k:"out"}
  ]},
  {name:"All records indexed?",k:"dec",children:[
    {name:"YES: QA phase",k:"out"},
    {name:"NO: debug logstash",k:"out"}
  ]}
]},
{name:"8. QA & Delivery",k:"step",children:[
  {name:"Products dashboard QA",k:"act"},
  {name:"Trends QA (after prod)",k:"act"},
  {name:"Kibana = Crawlboard?",k:"dec",children:[
    {name:"OK: share Kibana link",k:"out"},
    {name:"Fail: debug logstash",k:"out"}
  ]},
  {name:"QA passed?",k:"dec",children:[
    {name:"YES: deliver to client",k:"out"},
    {name:"NO: fix and re-QA",k:"out"}
  ]}
]}
]};
var NW=196,NH=48,NR=9;
var M={top:60,right:240,bottom:60,left:210};
var W=window.innerWidth-M.left-M.right;
var H=window.innerHeight-M.top-M.bottom;
var svg=d3.select("#tree").append("svg")
  .attr("width",W+M.left+M.right).attr("height",H+M.top+M.bottom)
  .call(d3.zoom().scaleExtent([.18,3]).on("zoom",e=>g.attr("transform",e.transform)))
  .on("dblclick.zoom",null);
var g=svg.append("g").attr("transform","translate("+M.left+","+(M.top+H/2)+")");
var i=0,dur=400;
var root=d3.hierarchy(data);
root.x0=0;root.y0=0;

var KF={start:"#1f2937",step:"#dbeafe",dec:"#fef3c7",out:"#dcfce7",act:"#f1f5f9"};
var KB={start:"#111827",step:"#3b82f6",dec:"#d97706",out:"#22c55e",act:"#94a3b8"};
var KT={start:"#fff",step:"#1e40af",dec:"#92400e",out:"#14532d",act:"#374151"};
function nFill(d){if(d.depth===0)return KF.start;return KF[d.data.k]||"#f3f4f6";}
function nBorder(d){if(d.depth===0)return KB.start;return KB[d.data.k]||"#d1d5db";}
function nText(d){if(d.depth===0)return KT.start;return KT[d.data.k]||"#374151";}
root.children.forEach(collapse);
update(root);
function collapse(d){if(d.children){d._children=d.children;d._children.forEach(collapse);d.children=null;}}
function wrap(t,n){if(t.length<=n)return[t];var m=t.lastIndexOf(" ",n);if(m<1)m=n;return[t.slice(0,m),t.slice(m+1)];}
function update(src){
  var tree=d3.tree().nodeSize([NH+28,1]);
  var td=tree(root);
  var nodes=td.descendants(),links=td.descendants().slice(1);
  var colW=Math.max(NW+36,W/5.4);nodes.forEach(d=>d.y=d.depth*colW);
  var node=g.selectAll("g.node").data(nodes,d=>d.id||(d.id=++i));
  var ne=node.enter().append("g").attr("class","node")
    .attr("transform",()=>"translate("+src.y0+","+src.x0+")")
    .on("click",click);
  ne.append("rect")
    .attr("width",NW).attr("height",NH).attr("x",-NW/2).attr("y",-NH/2)
    .attr("rx",NR).attr("ry",NR)
    .style("fill",d=>nFill(d)).style("stroke",d=>nBorder(d))
    .style("filter","drop-shadow(0 2px 5px rgba(0,0,0,.07))");
  ne.each(function(d){
    var el=d3.select(this),ln=wrap(d.data.name,23);
    var dy1=ln.length===1?"0.35em":"-0.5em", dy2="0.82em";
    el.append("text").attr("dy",dy1).attr("text-anchor","middle")
      .style("font-size","13px").style("fill",nText(d)).text(ln[0]);
    if(ln.length>1)
      el.append("text").attr("dy",dy2).attr("text-anchor","middle")
        .style("font-size","12px").style("fill",nText(d)).text(ln[1]);
  });
  ne.append("text").attr("class","ind")
    .attr("dy","0.35em").attr("x",NW/2-9).attr("text-anchor","middle")
    .style("font-size","9px");
  var nu=ne.merge(node);
  nu.transition().duration(dur).attr("transform",d=>"translate("+d.y+","+d.x+")");
  nu.select("rect").style("fill",d=>nFill(d)).style("stroke",d=>nBorder(d));
  nu.select(".ind")
    .text(d=>(d._children||d.children)?"●":"")
    .style("fill",d=>nBorder(d))
    .style("opacity",d=>(d._children||d.children)?1:0);
  node.exit().transition().duration(dur)
    .attr("transform",()=>"translate("+src.y+","+src.x+")").remove();
  var link=g.selectAll("path.link").data(links,d=>d.id);
  var le=link.enter().insert("path","g").attr("class","link")
    .attr("d",()=>{var o={x:src.x0,y:src.y0};return diag(o,o);})
    .style("stroke",d=>nBorder(d));
  le.merge(link).transition().duration(dur).attr("d",d=>diag(d,d.parent));
  link.exit().transition().duration(dur)
    .attr("d",()=>{var o={x:src.x,y:src.y};return diag(o,o);}).remove();
  nodes.forEach(d=>{d.x0=d.x;d.y0=d.y;});
}
function diag(s,d){
  return"M"+s.y+" "+s.x+" C"+(s.y+d.y)/2+" "+s.x+","+(s.y+d.y)/2+" "+d.x+","+d.y+" "+d.x;
}
function click(e,d){
  if(d.children){d._children=d.children;d.children=null;}
  else{d.children=d._children;d._children=null;}
  update(d);
}
</script>"""
    components.html(_html.replace("{_D3_INLINE}", _D3_INLINE), height=1000, scrolling=False)

def render_ops_map():
    page_title("Day-to-Day Operations Mind Map", "All 7 operational areas — expand any branch to explore tasks & tools.")
    st.markdown("""<div style="background:white;border:1px solid #e5e7eb;border-radius:10px;padding:10px 20px;margin-bottom:14px;font-size:0.82rem;color:#6b7280;box-shadow:0 1px 4px rgba(0,0,0,0.04);font-family:'Inter',sans-serif;">7 operational areas &mdash; expand any branch &nbsp;&middot;&nbsp; scroll to zoom &nbsp;&middot;&nbsp; drag to pan</div>""", unsafe_allow_html=True)
    _html = """<!DOCTYPE html><meta charset="utf-8">
<style>
html,body{margin:0;padding:0;width:100vw;height:100vh;overflow:hidden;
  background:#f0f2f6;font-family:'Inter','Segoe UI',-apple-system,sans-serif;}
#tree{width:100vw;height:100vh;}
.node rect{stroke-width:1.8px;transition:all .18s;cursor:pointer;}
.node rect:hover{opacity:.82;}
.node text{pointer-events:none;font-weight:500;}
.link{fill:none;stroke-width:1.6px;opacity:.4;}
#legend{position:absolute;top:12px;right:14px;background:rgba(255,255,255,.93);
  border:1px solid #e5e7eb;border-radius:9px;padding:10px 14px;
  font-size:12px;color:#374151;line-height:2.1;
  box-shadow:0 2px 8px rgba(0,0,0,.07);}
#hint{position:absolute;bottom:10px;right:14px;font-size:10px;color:#9ca3af;}
</style>
<div id="tree"></div>
<div id="legend"><b>Operational Areas</b><br><span style='color:#2563eb'>&#9646;</span> Kibana Monitoring<br><span style='color:#059669'>&#9646;</span> Input Sheets<br><span style='color:#d97706'>&#9646;</span> Cost Analysis<br><span style='color:#7c3aed'>&#9646;</span> Crawl Health<br><span style='color:#dc2626'>&#9646;</span> Mapping & Tracking<br><span style='color:#0891b2'>&#9646;</span> Maintenance<br><span style='color:#be185d'>&#9646;</span> Automation<br></div>
<div id="hint">Scroll = zoom &nbsp;·&nbsp; Drag = pan &nbsp;·&nbsp; Click = expand/collapse</div>
{_D3_INLINE}
<script>
var data={name:"Daily Operations Hub",children:[
{name:"Kibana Monitoring",children:[
  {name:"Client vs Site",children:[
    {name:"Active clients list"},
    {name:"Crawl frequency"},
    {name:"Records by site"},
    {name:"Client % share/site"}
  ]},
  {name:"Proxy Status",children:[
    {name:"Success/fail rate"},
    {name:"Oxylabs premium"},
    {name:"Weekly check"}
  ]},
  {name:"Disk Time Opt.",children:[
    {name:"Slow sites identify"},
    {name:"Infinity loop detect"},
    {name:"Weekly + on-demand"}
  ]},
  {name:"Extraction Duration",children:[
    {name:"Slow parsers"},
    {name:"XPath bottlenecks"},
    {name:"Weekly check"}
  ]},
  {name:"Cost Analytics",children:[
    {name:"Client vs Site dash"},
    {name:"Monthly infra cost"},
    {name:"% data per client"}
  ]}
]},
{name:"Input Sheet Mgmt",children:[
  {name:"Data Request Format",children:[
    {name:"Client URLs + pincodes"},
    {name:"Crawl scheduling"},
    {name:"pincode_uniq_id logic"},
    {name:"Closed clients view"},
    {name:"Kibana link viewer"}
  ]},
  {name:"Add / Update URLs",children:[
    {name:"Sheet-based input"},
    {name:"ES-based input"},
    {name:"Pincode CSV input"}
  ]},
  {name:"AppScript Automation",children:[
    {name:"Auto crawl tracking"},
    {name:"Threshold alerts"},
    {name:"Status updates"}
  ]}
]},
{name:"Cost Analysis",children:[
  {name:"Clientwise (Monthly)",children:[
    {name:"Current vs prev cycle"},
    {name:"n-month trend"},
    {name:"42s Clientwise sheet"}
  ]},
  {name:"Sitewise (Monthly)",children:[
    {name:"Cost per site"},
    {name:"% data per client"},
    {name:"42s Sitewise sheet"}
  ]},
  {name:"InfraCost Input",children:[
    {name:"Platform + DevOps"},
    {name:"Manual ES indexing"},
    {name:"42s input data sheet"}
  ]}
]},
{name:"Crawl Health",children:[
  {name:"Count Mismatch",children:[
    {name:"Kibana vs Crawlboard"},
    {name:"Threshold sheet"},
    {name:"42S avg threshold"}
  ]},
  {name:"Failure Logs",children:[
    {name:"lbhdf12_logstash.log"},
    {name:"lbhdf13_logstash.log"},
    {name:"Copied hourly to ex51"}
  ]},
  {name:"Misc + Dep Sites",children:[
    {name:"Misc processing tracker"},
    {name:"Dep data upload tracker"},
    {name:"Banner sites progress"}
  ]}
]},
{name:"Mapping & Tracking",children:[
  {name:"Client-Site Mapping",children:[
    {name:"Client to site view"},
    {name:"Site to client view"},
    {name:"Active sites list"}
  ]},
  {name:"Domain Mapping JSON",children:[
    {name:"Products to Trends link"},
    {name:"domain_numbered_sites"},
    {name:"Update on new site"}
  ]},
  {name:"42S Schema Sheet",children:[
    {name:"Field definitions"},
    {name:"Must-have fields"},
    {name:"Index-specific fields"}
  ]}
]},
{name:"Maintenance",children:[
  {name:"Weekly Tasks",children:[
    {name:"Disk time review"},
    {name:"Extraction duration"},
    {name:"Proxy health check"},
    {name:"Image count check"}
  ]},
  {name:"Monthly Tasks",children:[
    {name:"Retrospection doc"},
    {name:"Aggregated report"},
    {name:"Clientwise cost"},
    {name:"Sitewise cost"}
  ]},
  {name:"On-Demand Tasks",children:[
    {name:"New site addition"},
    {name:"Client pause/resume"},
    {name:"Threshold updates"},
    {name:"InfraCost update"}
  ]}
]},
{name:"Automation",children:[
  {name:"Google App Scripts",children:[
    {name:"Data Request Format"},
    {name:"Clientwise cost"},
    {name:"Sitewise cost"},
    {name:"Threshold alerts"}
  ]},
  {name:"Ruby Scripts",children:[
    {name:"Volume adjustment"},
    {name:"Missing upload check"},
    {name:"Weekly crawl tracker"}
  ]}
]}
]};
var NW=196,NH=48,NR=9;
var M={top:60,right:240,bottom:60,left:210};
var W=window.innerWidth-M.left-M.right;
var H=window.innerHeight-M.top-M.bottom;
var svg=d3.select("#tree").append("svg")
  .attr("width",W+M.left+M.right).attr("height",H+M.top+M.bottom)
  .call(d3.zoom().scaleExtent([.18,3]).on("zoom",e=>g.attr("transform",e.transform)))
  .on("dblclick.zoom",null);
var g=svg.append("g").attr("transform","translate("+M.left+","+(M.top+H/2)+")");
var i=0,dur=400;
var root=d3.hierarchy(data);
root.x0=0;root.y0=0;
var BC=["#2563eb", "#059669", "#d97706", "#7c3aed", "#dc2626", "#0891b2", "#be185d"];
function getBC(d){
  var a=d;while(a.depth>1&&a.parent)a=a.parent;
  if(a.depth===0)return"#1f2937";
  return BC[(a.parent?a.parent.children.indexOf(a):0)%BC.length];
}
function nFill(d){
  if(d.depth===0)return"#1f2937";
  if(d.depth===1){var c=getBC(d);return c+"18";}
  return"#f3f4f6";
}
function nBorder(d){return getBC(d);}
function nText(d){
  if(d.depth===0)return"#fff";
  if(d.depth===1)return getBC(d);
  return"#374151";
}
root.children.forEach(collapse);
update(root);
function collapse(d){if(d.children){d._children=d.children;d._children.forEach(collapse);d.children=null;}}
function wrap(t,n){if(t.length<=n)return[t];var m=t.lastIndexOf(" ",n);if(m<1)m=n;return[t.slice(0,m),t.slice(m+1)];}
function update(src){
  var tree=d3.tree().nodeSize([NH+28,1]);
  var td=tree(root);
  var nodes=td.descendants(),links=td.descendants().slice(1);
  var colW=Math.max(NW+36,W/5.4);nodes.forEach(d=>d.y=d.depth*colW);
  var node=g.selectAll("g.node").data(nodes,d=>d.id||(d.id=++i));
  var ne=node.enter().append("g").attr("class","node")
    .attr("transform",()=>"translate("+src.y0+","+src.x0+")")
    .on("click",click);
  ne.append("rect")
    .attr("width",NW).attr("height",NH).attr("x",-NW/2).attr("y",-NH/2)
    .attr("rx",NR).attr("ry",NR)
    .style("fill",d=>nFill(d)).style("stroke",d=>nBorder(d))
    .style("filter","drop-shadow(0 2px 5px rgba(0,0,0,.07))");
  ne.each(function(d){
    var el=d3.select(this),ln=wrap(d.data.name,23);
    var dy1=ln.length===1?"0.35em":"-0.5em", dy2="0.82em";
    el.append("text").attr("dy",dy1).attr("text-anchor","middle")
      .style("font-size","13px").style("fill",nText(d)).text(ln[0]);
    if(ln.length>1)
      el.append("text").attr("dy",dy2).attr("text-anchor","middle")
        .style("font-size","12px").style("fill",nText(d)).text(ln[1]);
  });
  ne.append("text").attr("class","ind")
    .attr("dy","0.35em").attr("x",NW/2-9).attr("text-anchor","middle")
    .style("font-size","9px");
  var nu=ne.merge(node);
  nu.transition().duration(dur).attr("transform",d=>"translate("+d.y+","+d.x+")");
  nu.select("rect").style("fill",d=>nFill(d)).style("stroke",d=>nBorder(d));
  nu.select(".ind")
    .text(d=>(d._children||d.children)?"●":"")
    .style("fill",d=>nBorder(d))
    .style("opacity",d=>(d._children||d.children)?1:0);
  node.exit().transition().duration(dur)
    .attr("transform",()=>"translate("+src.y+","+src.x+")").remove();
  var link=g.selectAll("path.link").data(links,d=>d.id);
  var le=link.enter().insert("path","g").attr("class","link")
    .attr("d",()=>{var o={x:src.x0,y:src.y0};return diag(o,o);})
    .style("stroke",d=>nBorder(d));
  le.merge(link).transition().duration(dur).attr("d",d=>diag(d,d.parent));
  link.exit().transition().duration(dur)
    .attr("d",()=>{var o={x:src.x,y:src.y};return diag(o,o);}).remove();
  nodes.forEach(d=>{d.x0=d.x;d.y0=d.y;});
}
function diag(s,d){
  return"M"+s.y+" "+s.x+" C"+(s.y+d.y)/2+" "+s.x+","+(s.y+d.y)/2+" "+d.x+","+d.y+" "+d.x;
}
function click(e,d){
  if(d.children){d._children=d.children;d.children=null;}
  else{d.children=d._children;d._children=null;}
  update(d);
}
</script>"""
    components.html(_html.replace("{_D3_INLINE}", _D3_INLINE), height=1000, scrolling=False)

def render_poc_guide():
    page_title("Task POC Guide", "Who to contact for every task type. Colour-coded by responsible team.")
    st.markdown("""<div style="background:white;border:1px solid #e5e7eb;border-radius:10px;padding:10px 20px;margin-bottom:14px;display:flex;gap:22px;flex-wrap:wrap;align-items:center;font-size:0.82rem;box-shadow:0 1px 4px rgba(0,0,0,0.04);font-family:'Inter',sans-serif;"><span style="color:#3b82f6;font-weight:600;">&#9646; Shanjai / Srinivas</span><span style="color:#7c3aed;font-weight:600;">&#9646; Dev Team</span><span style="color:#d97706;font-weight:600;">&#9646; Platform</span><span style="color:#dc2626;font-weight:600;">&#9646; TPM</span><span style="color:#16a34a;font-weight:600;">&#9646; DS / QA / Product</span></div>""", unsafe_allow_html=True)
    _html = """<!DOCTYPE html><meta charset="utf-8">
<style>
html,body{margin:0;padding:0;width:100vw;height:100vh;overflow:hidden;
  background:#f0f2f6;font-family:'Inter','Segoe UI',-apple-system,sans-serif;}
#tree{width:100vw;height:100vh;}
.node rect{stroke-width:1.8px;transition:all .18s;cursor:pointer;}
.node rect:hover{opacity:.82;}
.node text{pointer-events:none;font-weight:500;}
.link{fill:none;stroke-width:1.6px;opacity:.4;}
#legend{position:absolute;top:12px;right:14px;background:rgba(255,255,255,.93);
  border:1px solid #e5e7eb;border-radius:9px;padding:10px 14px;
  font-size:12px;color:#374151;line-height:2.1;
  box-shadow:0 2px 8px rgba(0,0,0,.07);}
#hint{position:absolute;bottom:10px;right:14px;font-size:10px;color:#9ca3af;}
</style>
<div id="tree"></div>
<div id="legend"><b>Point of Contact</b><br><span style='color:#3b82f6'>&#9646;</span> Shanjai / Srinivas<br><span style='color:#7c3aed'>&#9646;</span> Dev Team<br><span style='color:#d97706'>&#9646;</span> Platform Team<br><span style='color:#dc2626'>&#9646;</span> TPM<br><span style='color:#16a34a'>&#9646;</span> DS / QA / Product</div>
<div id="hint">Scroll = zoom &nbsp;·&nbsp; Drag = pan &nbsp;·&nbsp; Click = expand/collapse</div>
{_D3_INLINE}
<script>
var data={name:"Task POC Guide",k:"root",children:[
{name:"Site Setup",k:"dev",children:[
  {name:"New site needed",k:"tpm",children:[
    {name:"Contact: TPM",k:"tpm"},
    {name:"Provide: name + index"},
    {name:"PSS created by TPM"},
    {name:"Dev does setup after"}
  ]},
  {name:"Reuse existing site",k:"dev",children:[
    {name:"Contact: Dev team",k:"dev"},
    {name:"Add seed URLs only"},
    {name:"OK even if site active"}
  ]},
  {name:"Naming convention",k:"sh",children:[
    {name:"Contact: Shanjai",k:"sh"},
    {name:"Ref: 42S Documentation"},
    {name:"domain_tld{#}_{idx}_42s"}
  ]},
  {name:"Domain mapping JSON",k:"dev",children:[
    {name:"Contact: Dev team",k:"dev"},
    {name:"domain_numbered_sites"},
    {name:"_mapping.json"}
  ]}
]},
{name:"Schema & Fields",k:"plat",children:[
  {name:"New field addition",k:"plat",children:[
    {name:"Contact: Platform team",k:"plat"},
    {name:"Finalise name + type"},
    {name:"Platform updates DRL"},
    {name:"Then Dev implements"},
    {name:"Eg: weekly_units_sold"}
  ]},
  {name:"DRL / EXT changes",k:"dev",children:[
    {name:"Contact: Dev team",k:"dev"},
    {name:"After Platform mapping"}
  ]},
  {name:"Schema review",k:"sh",children:[
    {name:"Contact: Shanjai/Dev",k:"sh"},
    {name:"Check 42S schema sheet"}
  ]}
]},
{name:"Crawl Issues",k:"dev",children:[
  {name:"Crawl not running",k:"dev",children:[
    {name:"Contact: Dev → TPM",k:"dev"},
    {name:"Check crawl-board logs"},
    {name:"Check proxy status"}
  ]},
  {name:"Count mismatch",k:"sh",children:[
    {name:"Contact: Shanjai/Srinivas",k:"sh"},
    {name:"Kibana vs Crawlboard"},
    {name:"Check logstash logs"},
    {name:"Threshold sheet review"}
  ]},
  {name:"Proxy failures",k:"sh",children:[
    {name:"Contact: Srinivas/Shanjai",k:"sh"},
    {name:"proxy_overview dash"},
    {name:"Oxylabs premium check"}
  ]},
  {name:"Extraction errors",k:"dev",children:[
    {name:"Contact: Dev team",k:"dev"},
    {name:"XPath / JS page issues"},
    {name:"lbhdf12/13 logstash"}
  ]}
]},
{name:"Client Requirements",k:"sh",children:[
  {name:"New client intake",k:"sh",children:[
    {name:"Contact: Shanjai",k:"sh"},
    {name:"Classify seed URLs"},
    {name:"Feasibility check"},
    {name:"Coordinate with product"}
  ]},
  {name:"New Balance",k:"sh",children:[
    {name:"Image download: Shanjai",k:"sh"},
    {name:"Server upload: Platform",k:"plat"},
    {name:"Product matching: DS",k:"ds"},
    {name:"QA annotation: QA team",k:"ds"},
    {name:"New fields: Platform",k:"plat"}
  ]},
  {name:"RamyBrook",k:"dev",children:[
    {name:"JSON mapping: Dev",k:"dev"},
    {name:"Saks cart flow: Dev",k:"dev"},
    {name:"Python-Ruby: Dev",k:"dev"},
    {name:"Validation: DS+QA",k:"ds"}
  ]},
  {name:"Client escalation",k:"tpm",children:[
    {name:"Contact: TPM",k:"tpm"},
    {name:"TPM → Mgmt if needed"}
  ]}
]},
{name:"Cost & Infra",k:"plat",children:[
  {name:"Monthly cost report",k:"plat",children:[
    {name:"Contact: Platform+DevOps",k:"plat"},
    {name:"Generates infra spend"},
    {name:"Manually indexed to ES"}
  ]},
  {name:"Clientwise analysis",k:"sh",children:[
    {name:"Contact: Shanjai",k:"sh"},
    {name:"42s Clientwise sheet"},
    {name:"Current vs prev cycle"}
  ]},
  {name:"Sitewise analysis",k:"sh",children:[
    {name:"Contact: Shanjai",k:"sh"},
    {name:"42s Sitewise sheet"}
  ]},
  {name:"InfraCost input",k:"sh",children:[
    {name:"Contact: Shanjai/Srinivas",k:"sh"},
    {name:"42s input data sheet"}
  ]}
]},
{name:"Maintenance Tasks",k:"sh",children:[
  {name:"Weekly checks",k:"sh",children:[
    {name:"Contact: Shanjai/Srinivas",k:"sh"},
    {name:"Disk, Proxy, Extraction"},
    {name:"Image counts"}
  ]},
  {name:"Monthly reports",k:"sh",children:[
    {name:"Contact: Shanjai/Srinivas",k:"sh"},
    {name:"Retrospection doc"},
    {name:"Cost analysis sheets"}
  ]},
  {name:"On-demand updates",k:"sh",children:[
    {name:"Contact: Shanjai",k:"sh"},
    {name:"New site threshold"},
    {name:"Mapping sheet update"}
  ]}
]},
{name:"Escalation Path",k:"tpm",children:[
  {name:"Platform change",k:"plat",children:[
    {name:"Contact: Platform team",k:"plat"},
    {name:"Schema / DRL / infra"}
  ]},
  {name:"New site creation",k:"tpm",children:[
    {name:"Contact: TPM",k:"tpm"},
    {name:"PSS → Dev setup"}
  ]},
  {name:"Client SLA issue",k:"tpm",children:[
    {name:"Contact: TPM",k:"tpm"},
    {name:"TPM → Management"}
  ]}
]}
]};
var NW=196,NH=48,NR=9;
var M={top:60,right:240,bottom:60,left:210};
var W=window.innerWidth-M.left-M.right;
var H=window.innerHeight-M.top-M.bottom;
var svg=d3.select("#tree").append("svg")
  .attr("width",W+M.left+M.right).attr("height",H+M.top+M.bottom)
  .call(d3.zoom().scaleExtent([.18,3]).on("zoom",e=>g.attr("transform",e.transform)))
  .on("dblclick.zoom",null);
var g=svg.append("g").attr("transform","translate("+M.left+","+(M.top+H/2)+")");
var i=0,dur=400;
var root=d3.hierarchy(data);
root.x0=0;root.y0=0;

var KF={root:"#1f2937",sh:"#dbeafe",dev:"#ede9fe",plat:"#fef3c7",tpm:"#fee2e2",ds:"#dcfce7"};
var KB={root:"#111827",sh:"#3b82f6",dev:"#7c3aed",plat:"#d97706",tpm:"#dc2626",ds:"#16a34a"};
var KT={root:"#fff",sh:"#1e40af",dev:"#5b21b6",plat:"#92400e",tpm:"#991b1b",ds:"#14532d"};
function nFill(d){if(d.depth===0)return KF.root;return KF[d.data.k]||"#f3f4f6";}
function nBorder(d){if(d.depth===0)return KB.root;return KB[d.data.k]||"#d1d5db";}
function nText(d){if(d.depth===0)return KT.root;return KT[d.data.k]||"#374151";}
root.children.forEach(collapse);
update(root);
function collapse(d){if(d.children){d._children=d.children;d._children.forEach(collapse);d.children=null;}}
function wrap(t,n){if(t.length<=n)return[t];var m=t.lastIndexOf(" ",n);if(m<1)m=n;return[t.slice(0,m),t.slice(m+1)];}
function update(src){
  var tree=d3.tree().nodeSize([NH+28,1]);
  var td=tree(root);
  var nodes=td.descendants(),links=td.descendants().slice(1);
  var colW=Math.max(NW+36,W/5.4);nodes.forEach(d=>d.y=d.depth*colW);
  var node=g.selectAll("g.node").data(nodes,d=>d.id||(d.id=++i));
  var ne=node.enter().append("g").attr("class","node")
    .attr("transform",()=>"translate("+src.y0+","+src.x0+")")
    .on("click",click);
  ne.append("rect")
    .attr("width",NW).attr("height",NH).attr("x",-NW/2).attr("y",-NH/2)
    .attr("rx",NR).attr("ry",NR)
    .style("fill",d=>nFill(d)).style("stroke",d=>nBorder(d))
    .style("filter","drop-shadow(0 2px 5px rgba(0,0,0,.07))");
  ne.each(function(d){
    var el=d3.select(this),ln=wrap(d.data.name,23);
    var dy1=ln.length===1?"0.35em":"-0.5em", dy2="0.82em";
    el.append("text").attr("dy",dy1).attr("text-anchor","middle")
      .style("font-size","13px").style("fill",nText(d)).text(ln[0]);
    if(ln.length>1)
      el.append("text").attr("dy",dy2).attr("text-anchor","middle")
        .style("font-size","12px").style("fill",nText(d)).text(ln[1]);
  });
  ne.append("text").attr("class","ind")
    .attr("dy","0.35em").attr("x",NW/2-9).attr("text-anchor","middle")
    .style("font-size","9px");
  var nu=ne.merge(node);
  nu.transition().duration(dur).attr("transform",d=>"translate("+d.y+","+d.x+")");
  nu.select("rect").style("fill",d=>nFill(d)).style("stroke",d=>nBorder(d));
  nu.select(".ind")
    .text(d=>(d._children||d.children)?"●":"")
    .style("fill",d=>nBorder(d))
    .style("opacity",d=>(d._children||d.children)?1:0);
  node.exit().transition().duration(dur)
    .attr("transform",()=>"translate("+src.y+","+src.x+")").remove();
  var link=g.selectAll("path.link").data(links,d=>d.id);
  var le=link.enter().insert("path","g").attr("class","link")
    .attr("d",()=>{var o={x:src.x0,y:src.y0};return diag(o,o);})
    .style("stroke",d=>nBorder(d));
  le.merge(link).transition().duration(dur).attr("d",d=>diag(d,d.parent));
  link.exit().transition().duration(dur)
    .attr("d",()=>{var o={x:src.x,y:src.y};return diag(o,o);}).remove();
  nodes.forEach(d=>{d.x0=d.x;d.y0=d.y;});
}
function diag(s,d){
  return"M"+s.y+" "+s.x+" C"+(s.y+d.y)/2+" "+s.x+","+(s.y+d.y)/2+" "+d.x+","+d.y+" "+d.x;
}
function click(e,d){
  if(d.children){d._children=d.children;d.children=null;}
  else{d.children=d._children;d._children=null;}
  update(d);
}
</script>"""
    components.html(_html.replace("{_D3_INLINE}", _D3_INLINE), height=1000, scrolling=False)





# ─────────────────────────────────────────────────────────────────────────────
# COST CALCULATOR
# ─────────────────────────────────────────────────────────────────────────────

def _generate_cost_pdf(results, grand_total, selected_domains, platform_display, crawl_icons):
    """Build a formatted PDF cost estimate using ReportLab."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
    from reportlab.lib import pagesizes  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
    from reportlab.lib.enums import TA_CENTER  # type: ignore
    from reportlab.lib.colors import HexColor  # type: ignore
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=pagesizes.A4,
        topMargin=0.5*inch, bottomMargin=0.6*inch,
        leftMargin=0.6*inch, rightMargin=0.6*inch,
    )
    styles = getSampleStyleSheet()
    el = []

    title_s = ParagraphStyle("CCT", parent=styles["Heading1"], fontSize=18,
                              textColor=HexColor("#1f2937"), alignment=TA_CENTER,
                              fontName="Helvetica-Bold", spaceAfter=3)
    sub_s   = ParagraphStyle("CCS", parent=styles["Normal"], fontSize=9,
                              textColor=HexColor("#6b7280"), alignment=TA_CENTER, spaceAfter=12)
    sec_s   = ParagraphStyle("CCH", parent=styles["Normal"], fontSize=10,
                              textColor=HexColor("#ffffff"), backColor=HexColor("#1f2937"),
                              fontName="Helvetica-Bold", leftIndent=8, spaceBefore=10, spaceAfter=4)
    th_s    = ParagraphStyle("CCTH", parent=styles["Normal"], fontSize=7.5,
                              textColor=HexColor("#ffffff"), fontName="Helvetica-Bold")
    td_s    = ParagraphStyle("CCTD", parent=styles["Normal"], fontSize=8, textColor=HexColor("#111827"))
    tdr_s   = ParagraphStyle("CCTDR", parent=styles["Normal"], fontSize=8,
                              textColor=HexColor("#111827"), alignment=1)
    cost_s  = ParagraphStyle("CCCS", parent=styles["Normal"], fontSize=8,
                              textColor=HexColor("#dc2626"), fontName="Helvetica-Bold", alignment=2)
    note_s  = ParagraphStyle("CCN", parent=styles["Normal"], fontSize=7.5,
                              textColor=HexColor("#9ca3af"), alignment=TA_CENTER, spaceAfter=6)

    try:
        if os.path.exists(LOGO_PATH):
            logo = Image(LOGO_PATH, width=0.9*inch, height=0.72*inch)
            hdr_tbl = Table([[logo, Paragraph("<b>Crawl Cost Estimate</b>", title_s)]],
                            colWidths=[1.1*inch, 5.9*inch])
            hdr_tbl.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ]))
            el.append(hdr_tbl)
        else:
            el.append(Paragraph("<b>Crawl Cost Estimate</b>", title_s))
    except Exception:
        el.append(Paragraph("<b>Crawl Cost Estimate</b>", title_s))

    el.append(Paragraph(
        f"Generated: {date.today().strftime('%d %b %Y')} &nbsp;|&nbsp; "
        f"Grand Total: <b>${grand_total:,.4f}</b>",
        sub_s
    ))
    el.append(Spacer(1, 0.1*inch))

    col_widths = [1.8*inch, 0.85*inch, 0.6*inch, 0.5*inch, 0.65*inch, 0.8*inch, 0.8*inch]
    for domain in selected_domains:
        domain_results = [r for r in results if r["domain"] == domain]
        if not domain_results:
            continue
        display_name = platform_display.get(domain, domain)
        domain_total = sum(r["total_cost"] for r in domain_results)

        el.append(Paragraph(
            f"  {_html_mod.escape(display_name)} ({_html_mod.escape(domain)})", sec_s
        ))
        el.append(Spacer(1, 0.04*inch))

        header_row = [
            Paragraph("Crawl Type", th_s), Paragraph("Volume/Crawl", th_s),
            Paragraph("Freq", th_s),       Paragraph("Days", th_s),
            Paragraph("Zipcode", th_s),    Paragraph("Cost/Crawl", th_s),
            Paragraph("Total Cost", th_s),
        ]
        table_rows = [header_row]
        for r in domain_results:
            ct_icon = crawl_icons.get(r["crawl_type"], "")
            table_rows.append([
                Paragraph(_html_mod.escape(f"{ct_icon} {r['crawl_type']}"), td_s),
                Paragraph(f"{r['volume_per_crawl']:,}", tdr_s),
                Paragraph(f"{r['freq']}x/d", tdr_s),
                Paragraph(str(r["days"]), tdr_s),
                Paragraph(r["zip_mode"].replace(" Zipcode", ""), tdr_s),
                Paragraph(f"${r['cost_per_crawl']:,.4f}", cost_s),
                Paragraph(f"${r['total_cost']:,.4f}", cost_s),
            ])
        table_rows.append([
            Paragraph("<b>Subtotal</b>", td_s),
            Paragraph("", td_s), Paragraph("", td_s), Paragraph("", td_s),
            Paragraph("", td_s), Paragraph("", td_s),
            Paragraph(f"<b>${domain_total:,.4f}</b>", cost_s),
        ])

        t = Table(table_rows, colWidths=col_widths)
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  HexColor("#1f2937")),
            ("GRID",          (0, 0), (-1, -1), 0.4, HexColor("#e5e7eb")),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING",   (0, 0), (-1, -1), 8),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS",(0, 1), (-1, -2), [HexColor("#ffffff"), HexColor("#f9fafb")]),
            ("BACKGROUND",    (0, -1), (-1, -1), HexColor("#f1f5f9")),
        ]))
        el.append(t)
        el.append(Spacer(1, 0.15*inch))

    gt_data = [["", "", "", "", "", "",
                Paragraph(f"<b>Grand Total: ${grand_total:,.4f}</b>", cost_s)]]
    gt = Table(gt_data, colWidths=col_widths)
    gt.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), HexColor("#1f2937")),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
    ]))
    el.append(gt)
    el.append(Spacer(1, 0.12*inch))
    el.append(Paragraph(
        "Rates are benchmarks derived from internal crawl cost data. Actual costs may vary.",
        note_s
    ))
    doc.build(el)
    buffer.seek(0)
    return buffer.read()


def render_cost_calculator():
    page_title(
        "Cost Calculator",
        "Select platforms, configure crawl types per domain, and get a detailed cost estimate with PDF/CSV download."
    )

    # ── Load domain/rate config from CSV ──────────────────────────────────────
    _RATES_CSV = Path("crawl_cost_rates.csv")

    if not _RATES_CSV.exists():
        st.error("crawl_cost_rates.csv not found. Please add it next to app.py.")
        return

    PLATFORM_LIST, PLATFORM_DISPLAY, RATES = [], {}, {}
    try:
        with open(_RATES_CSV, newline="") as _fh:
            for _row in csv.DictReader(_fh):
                _d   = _row["domain"].strip()
                _zip = _row["zipcode"].strip().lower() == "true"
                _key = "with" if _zip else "without"
                if _d not in RATES:
                    PLATFORM_LIST.append(_d)
                    PLATFORM_DISPLAY[_d] = _row["display_name"].strip()
                    RATES[_d] = {}
                _sku = float(_row["sku_rate"])
                _cat = float(_row["cat_rate"])
                _kw  = float(_row["kw_rate"])
                if _sku < 0 or _cat < 0 or _kw < 0:
                    st.error(f"crawl_cost_rates.csv: negative rate for '{_d}' (zipcode={_row['zipcode']}). Rates must be ≥ 0.")
                    return
                RATES[_d][_key] = {"sku": _sku, "cat": _cat, "kw": _kw}
    except KeyError as e:
        st.error(f"crawl_cost_rates.csv is missing column: {e}. Expected columns: domain, display_name, zipcode, sku_rate, cat_rate, kw_rate")
        return
    except ValueError as e:
        st.error(f"crawl_cost_rates.csv contains an invalid number: {e}")
        return
    except Exception as e:
        st.error(f"Failed to load crawl_cost_rates.csv: {e}")
        return

    if not PLATFORM_LIST:
        st.error("crawl_cost_rates.csv loaded successfully but contains no domains. Add at least one domain row.")
        return

    # Remove stale selections: domains that were saved in session state but are
    # no longer present in the CSV (e.g. domain was renamed or removed).
    _saved_sel = st.session_state.get("cc_selected_domains", [])
    _stale = [d for d in _saved_sel if d not in PLATFORM_LIST]
    if _stale:
        st.warning(
            f"The following platform(s) are no longer in the rate config and have been removed from your selection: "
            + ", ".join(_stale)
        )
        st.session_state["cc_selected_domains"] = [d for d in _saved_sel if d in PLATFORM_LIST]

    CRAWL_TYPES = [
        "Category Based", "SKU / Product URL Based", "SOS (Share of Search)",
        "Reviews", "Keyword Level", "Festive Sales Day Crawl", "Banner Crawl",
    ]
    CRAWL_ICONS = {
        "Category Based": "🗂️", "SKU / Product URL Based": "📦",
        "SOS (Share of Search)": "🔍", "Reviews": "⭐",
        "Keyword Level": "🔑", "Festive Sales Day Crawl": "🎉", "Banner Crawl": "🖼️",
    }
    CRAWL_DESC = {
        "Category Based":          "Browse category pages with pagination",
        "SKU / Product URL Based": "Direct product URL / API fetch",
        "SOS (Share of Search)":   "Search result pages crawled by keyword",
        "Reviews":                 "Product review page crawling",
        "Keyword Level":           "Keyword-based search result crawling",
        "Festive Sales Day Crawl": "High-freq category crawl for sale events (1.2× rate)",
        "Banner Crawl":            "Promotional / banner URL monitoring ($0.001/URL/crawl)",
    }

    def get_rate(domain, crawl_type, with_zip):
        if domain not in RATES:
            return 0.0
        key = "with" if with_zip else "without"
        r = RATES[domain].get(key, RATES[domain].get("without", {}))
        if crawl_type == "Category Based":           return r.get("cat", 0.0)
        if crawl_type == "Festive Sales Day Crawl":  return r.get("cat", 0.0) * 1.2
        if crawl_type == "SKU / Product URL Based":  return r.get("sku", 0.0)
        if crawl_type == "Reviews":                  return r.get("sku", 0.0) * 0.7
        if crawl_type in ("SOS (Share of Search)", "Keyword Level"): return r.get("kw", 0.0)
        if crawl_type == "Banner Crawl":             return 0.001
        return 0.0

    def compute_volume(crawl_type, a, b):
        if crawl_type in ("Category Based", "Festive Sales Day Crawl"): return a * b
        if crawl_type == "SKU / Product URL Based":  return a
        if crawl_type == "Reviews":                  return a
        if crawl_type in ("SOS (Share of Search)", "Keyword Level"): return a * b
        if crawl_type == "Banner Crawl":             return a
        return 0

    def _cost_cell(val):
        if val == 0:      return '<span style="color:#16a34a;font-weight:700;">$0.00</span>'
        if val < 1:       return f'<span style="color:#ca8a04;font-weight:700;">${val:.4f}</span>'
        if val < 100:     return f'<span style="color:#ea580c;font-weight:700;">${val:,.4f}</span>'
        return f'<span style="color:#dc2626;font-weight:700;">${val:,.4f}</span>'

    # ── Step 1: Platform Selection ────────────────────────────────────────────
    section_header("🌐", "Step 1 — Select Platforms")
    selected_domains = st.multiselect(
        "Choose platforms to include in this estimate",
        options=PLATFORM_LIST,
        format_func=lambda x: PLATFORM_DISPLAY.get(x, x),
        key="cc_selected_domains",
        placeholder="Select one or more platforms...",
    )

    if not selected_domains:
        st.markdown("""
        <div style="text-align:center;padding:56px 20px;color:#94a3b8;font-family:'Inter',sans-serif;
        background:white;border-radius:14px;border:2px dashed #e5e7eb;margin-top:20px;">
            <div style="font-size:2.8rem;margin-bottom:14px;">📊</div>
            <div style="font-size:1rem;font-weight:600;color:#374151;">Select platforms above to begin</div>
            <div style="font-size:0.82rem;margin-top:6px;">
                Choose one or more platforms, configure crawl types for each,<br>then click Generate.
            </div>
        </div>""", unsafe_allow_html=True)
        return

    # ── Step 2: Per-Domain Configuration ─────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    section_header("⚙️", "Step 2 — Configure Crawl Types")

    for domain in selected_domains:
        display_name = PLATFORM_DISPLAY.get(domain, domain)

        with st.expander(f"**{display_name}**  ·  {domain}", expanded=True):
            col_ct, col_zip = st.columns([3, 1])
            with col_ct:
                selected_cts = st.multiselect(
                    "Crawl types",
                    options=CRAWL_TYPES,
                    format_func=lambda x: f"{CRAWL_ICONS.get(x, '')}  {x}",
                    key=f"cc_ct_{domain}",
                    placeholder="Select crawl type(s)...",
                )
            with col_zip:
                st.radio("Zipcode", ["Without Zipcode", "With Zipcode", "Both"],
                         key=f"cc_zip_{domain}")

            _zip_mode_now = st.session_state.get(f"cc_zip_{domain}", "Without Zipcode")
            if _zip_mode_now in ("With Zipcode", "Both"):
                _zc_col, _ = st.columns([1, 3])
                with _zc_col:
                    st.number_input("Number of Zipcodes", min_value=1, value=1, step=1,
                                    key=f"cc_zipcount_{domain}")

            if not selected_cts:
                st.caption("No crawl types selected for this platform.")
                continue

            for ct in selected_cts:
                st.markdown(
                    f'<div style="font-size:0.82rem;font-weight:600;color:#374151;'
                    f'margin:14px 0 6px 0;font-family:\'Inter\',sans-serif;">'
                    f'{CRAWL_ICONS.get(ct, "")} {ct}'
                    f'<span style="font-size:0.72rem;color:#9ca3af;font-weight:400;'
                    f'margin-left:8px;">— {CRAWL_DESC.get(ct, "")}</span></div>',
                    unsafe_allow_html=True,
                )
                if ct in ("Category Based", "Festive Sales Day Crawl"):
                    c1, c2, c3, c4 = st.columns(4)
                    with c1: st.number_input("Category URLs",   min_value=1, value=100, step=10, key=f"cc_{domain}_{ct}_a")
                    with c2: st.number_input("SKUs/Category",   min_value=1, value=50,  step=5,  key=f"cc_{domain}_{ct}_b")
                    with c3: st.number_input("Crawls/day",      min_value=1, value=1,   step=1,  key=f"cc_{domain}_{ct}_c")
                    with c4: st.number_input("Duration (days)", min_value=1, value=30,  step=1,  key=f"cc_{domain}_{ct}_d")
                elif ct == "SKU / Product URL Based":
                    c1, c2, c3 = st.columns(3)
                    with c1: st.number_input("Number of SKUs",  min_value=1, value=1000, step=100, key=f"cc_{domain}_{ct}_a")
                    with c2: st.number_input("Crawls/day",      min_value=1, value=1,    step=1,   key=f"cc_{domain}_{ct}_c")
                    with c3: st.number_input("Duration (days)", min_value=1, value=30,   step=1,   key=f"cc_{domain}_{ct}_d")
                elif ct == "Reviews":
                    c1, c2, c3 = st.columns(3)
                    with c1: st.number_input("Number of Products", min_value=1, value=500, step=50, key=f"cc_{domain}_{ct}_a")
                    with c2: st.number_input("Crawls/day",         min_value=1, value=1,   step=1,  key=f"cc_{domain}_{ct}_c")
                    with c3: st.number_input("Duration (days)",    min_value=1, value=30,  step=1,  key=f"cc_{domain}_{ct}_d")
                elif ct in ("SOS (Share of Search)", "Keyword Level"):
                    c1, c2, c3, c4 = st.columns(4)
                    with c1: st.number_input("Keywords",          min_value=1, value=200, step=10, key=f"cc_{domain}_{ct}_a")
                    with c2: st.number_input("SKUs/Keyword",      min_value=1, value=60,  step=5,  key=f"cc_{domain}_{ct}_b")
                    with c3: st.number_input("Crawls/day",        min_value=1, value=1,   step=1,  key=f"cc_{domain}_{ct}_c")
                    with c4: st.number_input("Duration (days)",   min_value=1, value=30,  step=1,  key=f"cc_{domain}_{ct}_d")
                elif ct == "Banner Crawl":
                    c1, c2, c3 = st.columns(3)
                    with c1: st.number_input("Banner URLs",       min_value=1, value=20,  step=5,  key=f"cc_{domain}_{ct}_a")
                    with c2: st.number_input("Crawls/day",        min_value=1, value=1,   step=1,  key=f"cc_{domain}_{ct}_c")
                    with c3: st.number_input("Duration (days)",   min_value=1, value=30,  step=1,  key=f"cc_{domain}_{ct}_d")
                st.markdown('<div style="height:1px;background:#f1f5f9;margin:6px 0 2px 0;"></div>',
                            unsafe_allow_html=True)

    # ── Generate button ───────────────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    _, btn_col, _ = st.columns([2, 1, 2])
    with btn_col:
        if st.button("📊  Generate Estimate", use_container_width=True, type="primary"):
            st.session_state["cc_show_results"] = True
            components.html(
                "<script>window.parent.document.querySelector('[data-testid=\"stAppViewContainer\"] > section')?.scrollTo({top:999999,behavior:'smooth'});</script>",
                height=0,
            )

    if not st.session_state.get("cc_show_results"):
        return

    # ── Compute results ───────────────────────────────────────────────────────
    results = []
    for domain in selected_domains:
        display_name = PLATFORM_DISPLAY.get(domain, domain)
        selected_cts = st.session_state.get(f"cc_ct_{domain}", [])
        zip_mode = st.session_state.get(f"cc_zip_{domain}", "Without Zipcode")
        zip_variants = (
            [("Without Zipcode", False), ("With Zipcode", True)]
            if zip_mode == "Both"
            else [(zip_mode, zip_mode == "With Zipcode")]
        )

        zip_count = st.session_state.get(f"cc_zipcount_{domain}", 1)
        for ct in selected_cts:
            a  = st.session_state.get(f"cc_{domain}_{ct}_a", 0)
            b  = st.session_state.get(f"cc_{domain}_{ct}_b", 0)
            c_ = st.session_state.get(f"cc_{domain}_{ct}_c", 1)
            d  = st.session_state.get(f"cc_{domain}_{ct}_d", 30)
            volume = compute_volume(ct, a, b)
            if volume == 0 and ct != "Banner Crawl":
                st.warning(f"**{display_name} — {ct}**: volume is 0. Check your inputs (e.g. number of SKUs / categories / keywords).")
            for zm, wz in zip_variants:
                effective_volume = volume * zip_count if wz else volume
                rate  = get_rate(domain, ct, wz)
                cpc   = effective_volume * rate
                total = cpc * c_ * d
                results.append({
                    "domain": domain, "display": display_name, "crawl_type": ct,
                    "volume_per_crawl": effective_volume, "freq": c_, "days": d,
                    "zip_mode": zm, "rate": rate,
                    "cost_per_crawl": cpc, "total_cost": total,
                })

    if not results:
        st.session_state["cc_show_results"] = False
        st.warning("No crawl types configured. Select crawl types for at least one platform.")
        return

    # ── Results header ────────────────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    section_header("📊", "Cost Estimate Results")

    grand_total = sum(r["total_cost"] for r in results)
    if grand_total == 0:
        st.info("All configured crawl types have a $0 rate. Check that the platforms and crawl types are correct, or update the rates in crawl_cost_rates.csv.")
    s1, s2, s3, s4 = st.columns(4)
    for col, lbl, val, accent in [
        (s1, "Grand Total (USD)", f"${grand_total:,.4f}", "#ef4444"),
        (s2, "Platforms",         str(len(set(r["domain"] for r in results))), "#1f2937"),
        (s3, "Crawl Configs",     str(len(results)),                           "#1f2937"),
        (s4, "Generated On",      datetime.now().strftime("%d %b %Y"),         "#1f2937"),
    ]:
        with col:
            st.markdown(f"""
            <div style="background:white;border-radius:12px;padding:16px 18px;
            border-left:4px solid {accent};box-shadow:0 2px 8px rgba(0,0,0,0.07);
            font-family:'Inter',sans-serif;">
                <div style="font-size:0.67rem;color:#94a3b8;text-transform:uppercase;
                letter-spacing:0.09em;font-weight:700;">{lbl}</div>
                <div style="font-size:1.15rem;font-weight:700;color:#0f172a;margin-top:5px;">{val}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Per-platform result tables ────────────────────────────────────────────
    for domain in selected_domains:
        domain_results = [r for r in results if r["domain"] == domain]
        if not domain_results:
            continue
        display_name = PLATFORM_DISPLAY.get(domain, domain)
        domain_total = sum(r["total_cost"] for r in domain_results)

        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#1f2937 0%,#374151 100%);
        border-radius:12px 12px 0 0;padding:12px 18px;display:flex;
        justify-content:space-between;align-items:center;font-family:'Inter',sans-serif;">
            <div style="font-size:0.95rem;font-weight:700;color:white;">
                {display_name}
                <span style="font-size:0.75rem;font-weight:400;color:#9ca3af;margin-left:6px;">({domain})</span>
            </div>
            <div style="font-size:0.9rem;font-weight:700;color:#fde68a;">
                Platform Total: ${domain_total:,.4f}
            </div>
        </div>""", unsafe_allow_html=True)

        rows_html = ""
        for i, r in enumerate(domain_results):
            bg   = "#ffffff" if i % 2 == 0 else "#f9fafb"
            icon = CRAWL_ICONS.get(r["crawl_type"], "")
            rows_html += (
                f'<tr style="background:{bg};border-bottom:1px solid #f1f5f9;">'
                f'<td style="padding:10px 16px;font-size:0.875rem;color:#0f172a;font-weight:500;">{icon} {r["crawl_type"]}</td>'
                f'<td style="padding:10px 16px;text-align:center;font-size:0.8rem;color:#374151;">{r["volume_per_crawl"]:,}</td>'
                f'<td style="padding:10px 16px;text-align:center;font-size:0.8rem;color:#374151;">{r["freq"]}×/day</td>'
                f'<td style="padding:10px 16px;text-align:center;font-size:0.8rem;color:#374151;">{r["days"]} days</td>'
                f'<td style="padding:10px 16px;text-align:center;font-size:0.75rem;color:#6b7280;">{r["zip_mode"].replace(" Zipcode","")}</td>'
                f'<td style="padding:10px 16px;text-align:right;">{_cost_cell(r["cost_per_crawl"])}</td>'
                f'<td style="padding:10px 16px;text-align:right;">{_cost_cell(r["total_cost"])}</td>'
                f'</tr>'
            )

        th = ("padding:9px 16px;font-size:0.7rem;text-transform:uppercase;"
              "letter-spacing:0.1em;color:#64748b;font-weight:700;background:#f8fafc;")
        st.markdown(f"""
        <div style="border:1px solid #e2e8f0;border-top:none;border-radius:0 0 12px 12px;
        overflow:hidden;margin-bottom:24px;box-shadow:0 4px 12px rgba(0,0,0,0.06);">
        <table style="width:100%;border-collapse:collapse;font-family:'Inter',sans-serif;">
        <thead><tr style="border-bottom:2px solid #e2e8f0;">
            <th style="{th}text-align:left;">Crawl Type</th>
            <th style="{th}text-align:center;">Volume/Crawl</th>
            <th style="{th}text-align:center;">Frequency</th>
            <th style="{th}text-align:center;">Duration</th>
            <th style="{th}text-align:center;">Zipcode</th>
            <th style="{th}text-align:right;">Cost/Crawl</th>
            <th style="{th}text-align:right;">Total Cost</th>
        </tr></thead>
        <tbody>{rows_html}</tbody>
        </table></div>""", unsafe_allow_html=True)

    # ── Downloads ─────────────────────────────────────────────────────────────
    section_header("📥", "Download Estimate")
    dl1, dl2, _ = st.columns([1, 1, 2])

    pdf_bytes = _generate_cost_pdf(results, grand_total, selected_domains, PLATFORM_DISPLAY, CRAWL_ICONS)
    with dl1:
        if st.download_button(
            "⬇️  Download PDF",
            data=pdf_bytes,
            file_name=f"cost_estimate_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
            mime="application/pdf",
            use_container_width=True,
        ):
            log_event(EVENT_DOWNLOAD_COST_PDF, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "cost_calc")

    csv_lines = ["Platform,Domain,Crawl Type,Volume/Crawl,Crawls/day,Days,Zipcode,Cost/Crawl (USD),Total Cost (USD)"]
    for r in results:
        csv_lines.append(
            f'{r["display"]},{r["domain"]},{r["crawl_type"]},'
            f'{r["volume_per_crawl"]},{r["freq"]},{r["days"]},{r["zip_mode"]},'
            f'{r["cost_per_crawl"]:.6f},{r["total_cost"]:.6f}'
        )
    csv_lines += ["", f'Grand Total,,,,,,,,{grand_total:.6f}']
    with dl2:
        if st.download_button(
            "⬇️  Download CSV",
            data="\n".join(csv_lines).encode(),
            file_name=f"cost_estimate_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv",
            use_container_width=True,
        ):
            log_event(EVENT_DOWNLOAD_COST_CSV, st.session_state.get("current_user", ""), st.session_state.get("analytics_sid", ""), "cost_calc")


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: ANALYTICS DASHBOARD  (admin only)
# ─────────────────────────────────────────────────────────────────────────────

def render_analytics():
    page_title("Analytics Dashboard", "Real-time usage insights across all users and features.")

    # ── Controls ───────────────────────────────────────────────────────────────
    period_options = {"Last 7 days": 7, "Last 30 days": 30, "Last 90 days": 90, "All time": 3650}
    ctrl1, ctrl2 = st.columns([3, 1])
    with ctrl1:
        period_label = st.selectbox("Time period", list(period_options.keys()), index=1, key="analytics_period")
    with ctrl2:
        st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
        if st.button("🔄  Refresh", use_container_width=True):
            st.rerun()
    days = period_options[period_label]
    summary = get_summary(days)

    # ── KPI cards ──────────────────────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)

    def _kpi(label, value, sub=""):
        sub_html = f'<div style="font-size:0.72rem;color:#9ca3af;margin-top:3px;">{sub}</div>' if sub else ""
        return f"""
        <div style="background:#fff;border:1px solid #e5e7eb;border-radius:10px;padding:16px 18px;
                    box-shadow:0 1px 4px rgba(0,0,0,0.05);font-family:'Inter',sans-serif;">
            <div style="font-size:0.72rem;font-weight:600;color:#9ca3af;text-transform:uppercase;
                        letter-spacing:0.08em;margin-bottom:6px;">{label}</div>
            <div style="font-size:1.7rem;font-weight:700;color:#1f2937;line-height:1;">{value}</div>
            {sub_html}
        </div>"""

    k1, k2, k3, k4, k5, k6, k7, k8 = st.columns(8)
    k1.markdown(_kpi("Sessions",        summary["total_sessions"],    f"Today: {summary['today_sessions']}"), unsafe_allow_html=True)
    k2.markdown(_kpi("Unique Users",    summary["unique_users"],      f"Today: {summary['today_users']}"),    unsafe_allow_html=True)
    k3.markdown(_kpi("Total Logins",    summary["login_count"]),                                              unsafe_allow_html=True)
    k4.markdown(_kpi("Total Events",    summary["total_events"]),                                             unsafe_allow_html=True)
    k5.markdown(_kpi("Docs Generated",  summary["docs_generated"],    "PDFs & feasibility"),                  unsafe_allow_html=True)
    k6.markdown(_kpi("Avg Pages/Session", summary["avg_session_depth"]),                                      unsafe_allow_html=True)
    k7.markdown(_kpi("Peak Hour",       summary["peak_hour_label"],   "Most active"),                         unsafe_allow_html=True)
    k8.markdown(_kpi("Top Page",        summary["most_visited"]),                                             unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Highlights strip ───────────────────────────────────────────────────────
    recent = summary["recent_events"]
    ua = summary["user_activity"]
    top_user = max(ua, key=lambda u: ua[u]) if ua else "—"
    pv = summary["page_views"]
    top_page = list(pv.keys())[0] if pv else "—"
    highlight_items = [
        ("🏆", "Most active user", top_user),
        ("📄", "Most visited page", top_page),
        ("⏰", "Peak usage hour", summary["peak_hour_label"]),
        ("📊", "Avg pages per session", str(summary["avg_session_depth"])),
        ("📝", "Docs / PDFs generated", str(summary["docs_generated"])),
    ]
    cols_h = st.columns(len(highlight_items))
    for col, (icon, label, val) in zip(cols_h, highlight_items):
        col.markdown(f"""
        <div style="background:linear-gradient(135deg,#f8fafc 0%,#f1f5f9 100%);
                    border:1px solid #e5e7eb;border-radius:8px;padding:12px 14px;
                    font-family:'Inter',sans-serif;text-align:center;">
            <div style="font-size:1.3rem;">{icon}</div>
            <div style="font-size:0.68rem;color:#9ca3af;text-transform:uppercase;
                        letter-spacing:0.08em;margin:4px 0 2px 0;font-weight:600;">{label}</div>
            <div style="font-size:0.88rem;font-weight:700;color:#1f2937;">{val}</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Activity over time ─────────────────────────────────────────────────────
    section_header("📈", "Activity Over Time")
    epd = summary["events_per_day"]
    if epd:
        df_epd = pd.DataFrame({"Date": list(epd.keys()), "Events": list(epd.values())}).set_index("Date")
        st.line_chart(df_epd, width="stretch", height=200)
    else:
        st.info("No activity recorded yet.")

    st.markdown("<br>", unsafe_allow_html=True)

    col_left, col_right = st.columns(2, gap="large")

    # ── Page views ────────────────────────────────────────────────────────────
    with col_left:
        section_header("📄", "Page Views by Section")
        if pv:
            df_pv = pd.DataFrame({"Page": list(pv.keys()), "Views": list(pv.values())}).set_index("Page")
            st.bar_chart(df_pv, width="stretch", height=220)
            # table with % share
            total_pv = sum(pv.values())
            rows_pv = [{"Page": k, "Views": v, "Share": f"{v/total_pv*100:.0f}%"} for k, v in pv.items()]
            st.dataframe(pd.DataFrame(rows_pv), width="stretch", hide_index=True)
        else:
            st.info("No page views recorded yet.")

    # ── Hourly heatmap ────────────────────────────────────────────────────────
    with col_right:
        section_header("🕐", "Hourly Activity Distribution")
        hd = summary["hourly_distribution"]
        df_hd = pd.DataFrame({"Hour": list(hd.keys()), "Events": list(hd.values())}).set_index("Hour")
        active_hours = df_hd[df_hd["Events"] > 0]
        if not active_hours.empty:
            st.bar_chart(active_hours, width="stretch", height=220)
            peak_h = active_hours["Events"].idxmax()
            st.caption(f"Peak hour: **{peak_h}** with **{active_hours.loc[peak_h, 'Events']}** events")
        else:
            st.info("No hourly data yet.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Per-user breakdown ────────────────────────────────────────────────────
    section_header("👥", "Per-User Breakdown")
    ub = summary["user_breakdown"]
    if ub:
        df_ub = pd.DataFrame(ub)
        st.dataframe(df_ub, width="stretch", hide_index=True)
    else:
        st.info("No user data yet.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Actions breakdown ─────────────────────────────────────────────────────
    col_a, col_b = st.columns(2, gap="large")
    with col_a:
        section_header("⚡", "Event Type Breakdown")
        actions = summary["actions"]
        if actions:
            total_acts = sum(actions.values())
            rows_act = [{"Event": k, "Count": v, "Share": f"{v/total_acts*100:.0f}%"} for k, v in actions.items()]
            st.dataframe(pd.DataFrame(rows_act), width="stretch", hide_index=True)
        else:
            st.info("No actions recorded yet.")

    with col_b:
        section_header("📊", "User Activity (Total Events)")
        if ua:
            df_ua = pd.DataFrame({"User": list(ua.keys()), "Events": list(ua.values())}).set_index("User")
            st.bar_chart(df_ua, width="stretch", height=220)
        else:
            st.info("No user activity yet.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Recent activity log ───────────────────────────────────────────────────
    section_header("🕐", "Recent Activity Log")
    if recent:
        rows = []
        for e in recent:
            ts_str = e.get("ts", "")
            try:
                ts_str = datetime.fromisoformat(ts_str).astimezone().strftime("%Y-%m-%d %H:%M:%S")
            except (ValueError, TypeError):
                pass
            rows.append({
                "Timestamp": ts_str,
                "User":      e.get("username", "—"),
                "Event":     EVENT_LABELS.get(e.get("event", ""), e.get("event", "—")),
                "Page":      PAGE_LABELS.get(e.get("page", ""), e.get("page", "—")) if e.get("page") else "—",
            })
        st.dataframe(pd.DataFrame(rows), width="stretch", hide_index=True)

        raw_events = load_events(days)
        if raw_events:
            csv_data = pd.DataFrame(raw_events).to_csv(index=False).encode()
            st.download_button(
                "⬇️  Export Full Log (CSV)",
                data=csv_data,
                file_name=f"analytics_log_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
            )
    else:
        st.info("No activity recorded yet.")


# ─────────────────────────────────────────────────────────────────────────────
# LOCALSTORAGE SESSION SYNC  (bidirectional component — no URL exposure)
# ─────────────────────────────────────────────────────────────────────────────
if st.session_state["ls_write_token"]:
    _tok = st.session_state["ls_write_token"]
    st.session_state["ls_write_token"] = ""
    _session_mgr(action="write", token=_tok, key="ls_write")

elif st.session_state["ls_clear"]:
    st.session_state["ls_clear"] = False
    _session_mgr(action="clear", token="", key="ls_clear")

elif not st.session_state["authenticated"]:
    # Read the stored token from localStorage.  Returns None on the first render
    # (before the component JS has responded), then the token string on the next.
    _stored_token = _session_mgr(action="read", token="", key="ls_read")
    if _stored_token:
        _sess_user, _sess_display = _load_session(_stored_token)
        if _sess_user:
            st.session_state["authenticated"] = True
            st.session_state["current_user"]  = _sess_user
            st.session_state["display_name"]  = _sess_display
            st.rerun()
        else:
            # Token is invalid/expired — wipe it from localStorage
            st.session_state["ls_clear"] = True
            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# ROUTER  (auth-gated)
# ─────────────────────────────────────────────────────────────────────────────

if not st.session_state["authenticated"]:
    render_login()
else:
    page = st.session_state["page"]
    # Log page_view once per actual navigation (not on every Streamlit rerun)
    if page != st.session_state.get("analytics_last_page", ""):
        log_event(
            EVENT_PAGE_VIEW,
            st.session_state.get("current_user", ""),
            st.session_state.get("analytics_sid", ""),
            page,
        )
        st.session_state["analytics_last_page"] = page
    if page == "main":
        render_main_form()
    elif page == "feasibility":
        render_feasibility()
    elif page == "req_flow":
        render_req_flow()
    elif page == "ops_map":
        render_ops_map()
    elif page == "poc_guide":
        render_poc_guide()
    elif page == "cost_calc":
        render_cost_calculator()
    elif page == "analytics":
        _role = (get_user(st.session_state.get("current_user", "") or "") or {}).get("role", "")
        if _role == "admin":
            render_analytics()
        else:
            st.error("Access denied. This page is restricted to administrators.")
            st.session_state["page"] = "main"
            st.rerun()
    elif page == "ext_tools":
        st.markdown("""
        <div style="font-family:'Inter',sans-serif;padding-bottom:8px;">
            <div style="font-size:1.2rem;font-weight:700;color:#1f2937;margin-bottom:4px;">External Tools &amp; Dashboards</div>
            <div style="font-size:0.82rem;color:#9ca3af;">Quick access to external platforms used by the team.</div>
        </div>
        <hr style="border:none;border-top:1px solid #e5e7eb;margin:12px 0 24px 0;">
        """, unsafe_allow_html=True)

        _EXT_TOOLS = [
            {
                "icon": "📓",
                "name": "NotebookLM",
                "desc": "AI-powered notebook for research, summarisation, and Q&A on uploaded documents.",
                "url": "https://notebooklm.google.com/notebook/1657537a-75b5-4f77-9228-24613e4d78ea?authuser=0&pli=1",
                "label": "Open NotebookLM",
            },
            {
                "icon": "📊",
                "name": "Kibana Dashboard",
                "desc": "Client vs Site — 42Signals Crawl Insights. Live crawl monitoring and analytics.",
                "url": "https://kibana42s-internal.promptcloud.com/app/dashboards#/view/e31a5fb0-41c3-11f0-ae05-5901704110bc?_g=(filters:!(),refreshInterval:(pause:!t,value:0),time:(from:'2025-08-19T18:30:00.000Z',to:now))",
                "label": "Open Kibana",
            },
            {
                "icon": "📋",
                "name": "Client Requirement Sheet",
                "desc": "Master Google Sheet tracking all client requirements and project details.",
                "url": "https://docs.google.com/spreadsheets/d/16vondEsN55P-HibdIGgrtxFTg02cc_VmiKU0pNJVR-8/edit?gid=1545244868#gid=1545244868",
                "label": "Open Sheet",
            },
            {
                "icon": "📉",
                "name": "42S Daily Threshold Sheet",
                "desc": "Daily threshold tracking sheet for 42Signals crawl volume and data pipeline.",
                "url": "https://docs.google.com/spreadsheets/d/1pDjioZXBY0TtFBb-sSQvzMXbWdBl_6xwTlCHrWbrQAY/edit?gid=0#gid=0",
                "label": "Open Sheet",
            },
            {
                "icon": "🎯",
                "name": "Redmine Agile Board",
                "desc": "Team ticket tracker — active sprints, assignments, and burndown chart for the 42S project.",
                "url": "https://redmine.promptcloud.com/agile/board?set_filter=1&f%5B%5D=project_id&op%5Bproject_id%5D=%3D&v%5Bproject_id%5D%5B%5D=40733&f%5B%5D=assigned_to_id&op%5Bassigned_to_id%5D=%21&v%5Bassigned_to_id%5D%5B%5D=57565&v%5Bassigned_to_id%5D%5B%5D=97098&v%5Bassigned_to_id%5D%5B%5D=77277&v%5Bassigned_to_id%5D%5B%5D=5972&v%5Bassigned_to_id%5D%5B%5D=60148&v%5Bassigned_to_id%5D%5B%5D=96645&v%5Bassigned_to_id%5D%5B%5D=97240&v%5Bassigned_to_id%5D%5B%5D=96993&v%5Bassigned_to_id%5D%5B%5D=53918&v%5Bassigned_to_id%5D%5B%5D=97369&v%5Bassigned_to_id%5D%5B%5D=95248&v%5Bassigned_to_id%5D%5B%5D=85453&v%5Bassigned_to_id%5D%5B%5D=96640&v%5Bassigned_to_id%5D%5B%5D=84753&v%5Bassigned_to_id%5D%5B%5D=86398&f%5B%5D=status_id&op%5Bstatus_id%5D=%3D&f_status%5B%5D=1&f_status%5B%5D=8&f_status%5B%5D=7&f_status%5B%5D=21&f_status%5B%5D=20&f_status%5B%5D=2&f_status%5B%5D=25&f_status%5B%5D=74&f_status%5B%5D=4&c%5B%5D=day_in_state&c%5B%5D=parent&default_chart=burndown_chart&chart_unit=issues&group_by=assigned_to&color_base=priority",
                "label": "Open Redmine",
            },
            {
                "icon": "✍️",
                "name": "Redmine Ticket Creation",
                "desc": "To create any new tickets",
                "url": "https://redmine.promptcloud.com/issues/new",
                "label": "Open Redmine",
            },
        ]

        cols = st.columns(2, gap="large")
        for i, tool in enumerate(_EXT_TOOLS):
            with cols[i % 2]:
                st.markdown(f"""
                <div style="
                    background:#fff;
                    border:1px solid #e5e7eb;
                    border-radius:12px;
                    padding:20px 22px;
                    margin-bottom:16px;
                    box-shadow:0 1px 4px rgba(0,0,0,0.05);
                    font-family:'Inter',sans-serif;
                ">
                    <div style="font-size:1.6rem;margin-bottom:10px;">{tool["icon"]}</div>
                    <div style="font-size:0.95rem;font-weight:700;color:#1f2937;margin-bottom:6px;">{tool["name"]}</div>
                    <div style="font-size:0.8rem;color:#6b7280;line-height:1.55;margin-bottom:16px;">{tool["desc"]}</div>
                    <a href="{tool["url"]}" target="_blank" style="
                        display:inline-block;
                        background:linear-gradient(135deg,#1f2937 0%,#374151 100%);
                        color:#fff;
                        text-decoration:none;
                        font-size:0.78rem;
                        font-weight:600;
                        padding:8px 18px;
                        border-radius:6px;
                        letter-spacing:0.02em;
                        box-shadow:0 1px 4px rgba(0,0,0,0.12);
                    ">{tool["label"]} ↗</a>
                </div>
                """, unsafe_allow_html=True)
