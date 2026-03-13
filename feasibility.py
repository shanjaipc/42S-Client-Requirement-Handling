import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO
import os

st.title("Crawl Feasibility Request Form")

# Initialize session state for domain inputs
if "domain_count" not in st.session_state:
    st.session_state.domain_count = 1

st.subheader("Requirement Information")

client_name = st.text_input("Client Name")
requestor_name = st.text_input("Requestor Name")

st.subheader("Domain List")

num_domains = st.number_input(
    "Number of Domains",
    min_value=1,
    step=1,
    value=1
)

domains = []

for i in range(int(num_domains)):
    domain = st.text_input(
        f"Domain {i+1}",
        placeholder="example.com",
        key=f"domain_{i}"
    )
    if domain:
        domains.append(domain)

domain_list = "\n".join(domains)

st.subheader("Crawl Configuration")

crawl_options = st.multiselect(
    "Select Crawl Type and Special Requirements",
    [
        "Category Based",
        "Product URL Input Based",
        "SOS",
        "Reviews",
        "Festive Sales Day Crawl",
        "Banner Crawl",
        "Others"
    ]
)

crawl_type = None
if "Category Based" in crawl_options:
    crawl_type = "Category Based"
elif "Product URL Input Based" in crawl_options:
    crawl_type = "Product URL Input Based"

crawl_features = [
    opt for opt in crawl_options
    if opt not in ["Category Based", "Product URL Input Based"]
]

others_desc = ""
if "Others" in crawl_features:
    others_desc = st.text_input("If Others, specify")

st.subheader("Zipcode Requirement")

zipcode_type = st.radio(
    "Zipcode Handling",
    ["With Zipcode", "Without Zipcode", "Both"]
)

# Target location appears dynamically
target_city = ""
target_state = ""
target_country = ""

if zipcode_type in ["With Zipcode", "Both"]:
    st.subheader("Target Location")

    target_city = st.text_input("City")
    target_state = st.text_input("State")
    target_country = st.text_input("Country")

st.subheader("Additional Information")

additional_notes = st.text_area("Additional Details")

if st.button("Generate Document"):

    document = Document()

    document.add_heading(f"{client_name} Feasibility Requirement", level=1)

    document.add_heading("Requirement Information", level=2)
    document.add_paragraph(f"Client Name: {client_name}")
    document.add_paragraph(f"Requestor Name: {requestor_name}")

    document.add_heading("Domains", level=2)
    for d in domains:
        document.add_paragraph(d)

    document.add_heading("Crawl Configuration", level=2)
    document.add_paragraph(f"Crawl Type: {crawl_type}")
    document.add_paragraph(f"Special Requirements: {', '.join(crawl_features)}")

    if others_desc:
        document.add_paragraph(f"Others Description: {others_desc}")

    document.add_heading("Zipcode Requirement", level=2)
    document.add_paragraph(f"Zipcode Handling: {zipcode_type}")

    if zipcode_type in ["With Zipcode", "Both"]:
        document.add_heading("Target Location", level=2)
        document.add_paragraph(f"City: {target_city}")
        document.add_paragraph(f"State: {target_state}")
        document.add_paragraph(f"Country: {target_country}")

    document.add_heading("Additional Notes", level=2)
    document.add_paragraph(additional_notes)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Request Document",
        data=buffer,
        file_name=f"{client_name}_feasibility_requirement.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )